"""
Tests for the APA 7 writing lane.

This module exists because of a complaint, not a bug report: asked for a place
to write in APA 7, the application offered a screen of rules and an export
button at the end of an eight-step research pipeline. Reading the rules is not
writing, and a formatter reachable only by first inventing a systematic review
is a formatter you do not have.

So the checks here are about the two things that lane has to get right. The
markup has to become the right document elements — a Level 4 heading that runs
into its paragraph, a forty-word quotation that becomes a block without its
quotation marks. And the findings have to fire when a rule is broken and stay
quiet when it is not, because a checker that cries wolf is one people learn to
export past.

Every finding is tested in both directions for that reason: the draft that must
be flagged, and the correct draft that must not be.

Run: python3 tests/test_compose.py
"""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from app.apa import compose                                   # noqa: E402
from app.apa.citations import plain                           # noqa: E402

FAILURES: list[str] = []
CHECKS = 0


def check(label: str, got, want) -> None:
    global CHECKS
    CHECKS += 1
    if got != want:
        FAILURES.append(f"{label}\n     got: {got!r}\n    want: {want!r}")


def rules(findings) -> list[str]:
    return [f["rule"] for f in findings]


def sound(**over) -> compose.Paper:
    """A paper with nothing wrong with it, to vary one thing at a time.

    Every check below that asserts a finding *fires* also has this as its
    control: without one, a checker that returned the same finding for every
    input would pass the entire suite.
    """
    paper = compose.Paper(
        paper_id="t", title="Nurse Staffing and Falls", authors=["Koch, M."],
        affiliations=["School of Nursing"], course="NURS 601",
        instructor="Dr. Reyes", due_date="September 1, 2026",
        body="# Introduction\n\nFalls are common on medical-surgical units.",
    )
    for name, value in over.items():
        setattr(paper, name, value)
    return paper


# ================================================================= the markup


def test_the_markup_becomes_the_right_elements() -> None:
    blocks = compose.parse_body(
        "# One\n\n## Two\n\n### Three\n\n#### Four\n\n##### Five\n\n"
        "A paragraph.\n\n- An item\n\n> A short quote"
    )
    kinds = [(b.kind, b.level, b.text or plain(b.runs)) for b in blocks]
    check("five heading levels", [k[1] for k in kinds if k[0] == "heading"],
          [1, 2, 3, 4, 5])
    check("a paragraph", ("paragraph", 0, "A paragraph.") in kinds, True)
    check("a list item", ("list-item", 0, "An item") in kinds, True)
    check("a short quote when marked", ("quote", 0, "A short quote") in kinds, True)


def test_a_soft_wrap_is_not_a_new_paragraph() -> None:
    """Pasting from a PDF arrives one line at a time. Treating each as its own
    paragraph would indent every line and double-space the wreckage."""
    blocks = compose.parse_body("One sentence\nrunning over\nthree lines.\n\nNext.")
    check("two paragraphs", len(blocks), 2)
    check("joined", plain(blocks[0].runs), "One sentence running over three lines.")


def test_a_long_quotation_becomes_a_block_without_its_marks() -> None:
    """APA 7 §8.27. Forty words or more is an indented block, and keeping the
    quotation marks as well is the commonest block-quote error."""
    quoted = " ".join(["word"] * 45)
    blocks = compose.parse_body(f'“{quoted}” (Smith, 2020, p. 4)')
    check("one block", len(blocks), 1)
    check("set as a quote", blocks[0].kind, "quote")
    check("marks removed", "“" in plain(blocks[0].runs), False)
    check("the words survive", plain(blocks[0].runs).startswith("word word"), True)

    short = " ".join(["word"] * 20)
    blocks = compose.parse_body(f'“{short}” (Smith, 2020, p. 4)')
    check("a short quotation stays in the paragraph", blocks[0].kind, "paragraph")
    check("and keeps its marks", "“" in plain(blocks[0].runs), True)


def test_asterisks_become_italics() -> None:
    blocks = compose.parse_body("The term *incidence* is defined here.")
    runs = blocks[0].runs
    check("three runs", len(runs), 3)
    check("the middle one is italic", [r.italic for r in runs], [False, True, False])
    check("and the asterisks are gone", plain(runs),
          "The term incidence is defined here.")


def test_the_outline_reports_shape_and_words() -> None:
    blocks = compose.parse_body("# One\n\nfour words go here\n\n## Two\n\ntwo words")
    check("outline", compose.outline(blocks),
          [{"level": 1, "text": "One", "words": 4},
           {"level": 2, "text": "Two", "words": 2}])
    check("words exclude headings", compose.word_count(blocks), 6)


# ================================================================= the checks


def test_a_sound_paper_reports_nothing() -> None:
    """The control for everything below. A checker that flagged every draft
    would satisfy each 'it fires' test and be worthless."""
    check("nothing to report", compose.check(sound()), [])


def test_a_skipped_heading_level_is_caught() -> None:
    """APA 7 §2.27. The structural error nobody catches by re-reading, because
    the page looks perfectly tidy."""
    findings = compose.check(sound(body="# One\n\ntext\n\n### Three\n\nmore"))
    check("flagged", "§2.27" in rules(findings), True)
    check("as an error", findings[0]["severity"], "error")
    check("naming the heading", "Three" in findings[0]["message"], True)

    check("but not when the levels are in order",
          rules(compose.check(sound(
              body="# One\n\ntext\n\n## Two\n\ntext\n\n### Three\n\ntext"))), [])


def test_the_title_page_must_carry_what_its_variant_requires() -> None:
    check("a student paper needs the course",
          "§2.7" in rules(compose.check(sound(course=""))), True)
    check("and the due date",
          "§2.8" in rules(compose.check(sound(due_date=""))), True)
    check("and must not carry a running head",
          "§2.4" in rules(compose.check(sound(running_head="SOMETHING"))), True)

    professional = sound(variant="professional", running_head="STAFFING AND FALLS",
                         course="", instructor="", due_date="")
    check("a professional paper is content without them",
          compose.check(professional), [])
    check("but not with an over-long running head",
          "§2.4" in rules(compose.check(sound(
              variant="professional", running_head="A" * 51,
              course="", instructor="", due_date=""))), True)
    check("nor a lower-case one",
          "§2.4" in rules(compose.check(sound(
              variant="professional", running_head="Staffing and Falls",
              course="", instructor="", due_date=""))), True)
    check("exactly fifty characters is allowed",
          "§2.4" in rules(compose.check(sound(
              variant="professional", running_head="A" * 50,
              course="", instructor="", due_date=""))), False)


def test_every_reference_must_be_cited_and_every_citation_referenced() -> None:
    """APA 7 §8.1 and §9.51 — the rule that costs marks. Checked in both
    directions, and asymmetrically on purpose: going from a reference to the
    text is a fact, going the other way is a guess."""
    work = compose.work_from_fields({
        "authors": "Aiken, Linda H.\nSloane, Douglas M.\nGriffiths, Peter",
        "year": "2021", "title": "Nursing skill mix and patient outcomes",
        "container": "BMJ Quality & Safety", "volume": "30", "pages": "639-647"})

    uncited = compose.check(sound(references=[work]))
    check("a reference nobody cited is an error", "§9.51" in rules(uncited), True)
    check("stated as an error", uncited[0]["severity"], "error")

    cited = sound(references=[work], body=(
        "# Introduction\n\nStaffing is associated with harm "
        "(Aiken et al., 2021)."))
    check("and silent once it is cited", compose.check(cited), [])

    check("a quotation locator does not break the match",
          compose.check(sound(references=[work], body=(
              '# Introduction\n\nThey found "a measurable association" '
              '(Aiken et al., 2021, p. 641).'))), [])

    check("the narrative form counts too",
          compose.check(sound(references=[work], body=(
              "# Introduction\n\nAiken et al. (2021) found an association."))), [])

    orphan = compose.check(sound(body=(
        "# Introduction\n\nAs Garcia (2018) showed, things happen.")))
    check("a citation with no reference is raised", "§8.1" in rules(orphan), True)
    check("but only as something to look at, never as an assertion",
          orphan[0]["severity"], "check")


def test_a_citation_in_the_wrong_form_is_raised() -> None:
    """Three or more authors take et al. from the *first* citation in APA 7 —
    the change from APA 6 that catches everyone."""
    work = compose.work_from_fields({
        "authors": "Aiken, Linda H.\nSloane, Douglas M.\nGriffiths, Peter",
        "year": "2021", "title": "Nursing skill mix", "container": "BMJ"})
    apa6 = compose.check(sound(references=[work], body=(
        "# Introduction\n\nHarm rises (Aiken, Sloane, & Griffiths, 2021).")))
    check("raised", "§8.17" in rules(apa6), True)
    check("with the correct form in the message",
          "(Aiken et al., 2021)" in
          next(f for f in apa6 if f["rule"] == "§8.17")["message"], True)


def test_one_source_listed_twice_is_caught() -> None:
    fields = {"authors": "Aiken, L.", "year": "2021",
              "title": "Nursing skill mix and patient outcomes"}
    first = compose.work_from_fields(dict(fields, doi="10.1/x"))
    second = compose.work_from_fields(fields)
    body = "# Introduction\n\nHarm rises (Aiken, 2021)."
    check("flagged", "§9.51" in rules(compose.check(
        sound(references=[first, second], body=body))), True)
    check("and one entry alone is fine",
          "§9.51" in rules(compose.check(
              sound(references=[first], body=body))), False)


def test_a_direct_quotation_needs_a_locator() -> None:
    """APA 7 §8.13."""
    long_quote = "> " + " ".join(["word"] * 45)
    check("a block quote without one is an error",
          "§8.13" in rules(compose.check(sound(body=long_quote))), True)
    check("and with one it is fine",
          "§8.13" in rules(compose.check(
              sound(body=long_quote + " (Smith, 2020, p. 4)"))), False)


def test_a_forty_word_quotation_left_inline_is_an_error() -> None:
    inline = 'They wrote "' + " ".join(["word"] * 45) + '" (Smith, 2020, p. 4).'
    findings = compose.check(sound(body=f"# Introduction\n\n{inline}"))
    check("flagged", "§8.27" in rules(findings), True)
    check("with the word count", "45 words" in findings[0]["message"], True)


def test_the_abstract_band_is_guidance_not_a_rule() -> None:
    check("too long is a warning, not an error",
          [f["severity"] for f in compose.check(sound(abstract=" ".join(["w"] * 300)))
           if f["rule"] == "§2.9"], ["warn"])
    check("and 200 words says nothing",
          "§2.9" in rules(compose.check(sound(abstract=" ".join(["w"] * 200)))), False)


def test_findings_are_ordered_worst_first() -> None:
    paper = sound(body="# One\n\ntext\n\n### Three\n\nAs Garcia (2018) showed.",
                  abstract=" ".join(["w"] * 300))
    severities = [f["severity"] for f in compose.check(paper)]
    check("errors, then warnings, then checks", severities,
          sorted(severities, key=lambda s: {"error": 0, "warn": 1}.get(s, 2)))


# =============================================================== the reference


def test_a_reference_is_built_and_rendered_by_the_shared_engine() -> None:
    """The formatting must be the same formatting the research lane uses, or a
    rule fixed in one place stays broken in the other."""
    work = compose.work_from_fields({
        "work_type": "journal-article",
        "authors": "Aiken, Linda H.\nSloane, Douglas M.\nGriffiths, Peter",
        "year": "2021", "title": "Nursing Skill Mix And Patient Outcomes",
        "container": "BMJ Quality & Safety", "volume": "30", "issue": "8",
        "pages": "639-647", "doi": "10.1136/bmjqs-2020-011512"})
    row = compose.reference_preview([work])[0]
    check("the reference entry", row["reference"],
          "Aiken, L. H., Sloane, D. M., & Griffiths, P. (2021). Nursing skill "
          "mix and patient outcomes. BMJ Quality & Safety, 30(8), 639–647. "
          "https://doi.org/10.1136/bmjqs-2020-011512")
    check("the parenthetical form", row["parenthetical"], "(Aiken et al., 2021)")
    check("the narrative form", row["narrative"], "Aiken et al. (2021)")


def test_an_organisation_is_kept_whole() -> None:
    """Mangling "World Health Organization" into "Organization, W. H." is the
    worst failure a reference builder has."""
    work = compose.work_from_fields({
        "work_type": "report", "authors": "World Health Organization",
        "year": "2023", "title": "Global report on infection prevention",
        "publisher": "World Health Organization"})
    row = compose.reference_preview([work])[0]
    check("whole in the reference",
          row["reference"].startswith("World Health Organization. (2023)."), True)
    check("and whole in text", row["parenthetical"],
          "(World Health Organization, 2023)")


# =================================================================== the store


def test_a_paper_survives_a_round_trip() -> None:
    with tempfile.TemporaryDirectory() as directory:
        store = compose.PaperStore(Path(directory))
        paper = store.new_paper("Staffing and Falls", "professional")
        paper.authors = ["Koch, M."]
        paper.body = "# One\n\n*Italic* text."
        paper.references = [compose.work_from_fields({
            "authors": "Aiken, L.", "year": "2021", "title": "A title"})]
        store.save(paper)

        loaded = store.load(paper.paper_id)
        check("title", loaded.title, "Staffing and Falls")
        check("variant", loaded.variant, "professional")
        check("a running head was derived", bool(loaded.running_head), True)
        check("authors", loaded.authors, ["Koch, M."])
        check("body", loaded.body, paper.body)
        check("references survive as Works", len(loaded.references), 1)
        check("with their authors parsed",
              loaded.references[0].authors[0].family, "Aiken")
        check("and their type", loaded.references[0].work_type.value,
              "journal-article")

        check("listed", [row["paper_id"] for row in store.list_papers()],
              [paper.paper_id])
        check("deleted", store.delete(paper.paper_id), True)
        check("and gone", store.list_papers(), [])


def test_a_paper_id_cannot_escape_its_directory() -> None:
    with tempfile.TemporaryDirectory() as directory:
        store = compose.PaperStore(Path(directory))
        for bad in ("../etc/passwd", "..", ".", "/etc/passwd"):
            try:
                store.path_for(bad)
                inside = False
            except ValueError:
                inside = True
            else:
                inside = Path(directory) in store.path_for(bad).parents
            check(f"{bad!r} stays inside the store", inside, True)


# ================================================================== the export


def test_the_document_is_apa_7_in_the_file() -> None:
    """Asserted against the saved OOXML rather than against our own intent.

    The run-in heading is the one this exists for: APA 7 §2.27 makes Levels 4
    and 5 continue on the same line as their paragraph, and the generic body
    builder emits every heading as its own paragraph — right for 1-3, wrong
    for 4 and 5.
    """
    import docx
    from app.apa.citations import CitationContext
    from app.apa.document import ApaPaper

    paper = sound(
        abstract="An abstract.", keywords=["staffing"],
        body=("# Introduction\n\nA paragraph of text.\n\n"
              "## Background\n\n### Prior Work\n\nMore text.\n\n"
              "#### Appraisal Method\n\nEach study was appraised.\n\n"
              "> A quoted passage (Aiken, 2021, p. 4)"),
        references=[compose.work_from_fields({
            "authors": "Aiken, L.", "year": "2021", "title": "A title",
            "container": "BMJ"})])

    blocks = compose.parse_body(paper.body)
    builder = ApaPaper(paper.as_project(), CitationContext(paper.references),
                       font="Times New Roman")
    builder.title_page()
    builder.abstract(paper.abstract, paper.keywords)
    compose.write_body(builder, blocks, paper.title)
    builder.references(paper.references)

    with tempfile.TemporaryDirectory() as directory:
        path = Path(directory) / "paper.docx"
        builder.save(str(path))
        document = docx.Document(str(path))
        paragraphs = [p.text for p in document.paragraphs]
        texts = "\n".join(paragraphs)

        check("the title page carries the title",
              paper.title in texts, True)
        check("and the student fields",
              all(v in texts for v in (paper.course, paper.instructor,
                                       paper.due_date)), True)
        check("the abstract has its own heading", "Abstract" in paragraphs, True)
        check("the Level 4 heading runs into its paragraph",
              "Appraisal Method. Each study was appraised." in paragraphs, True)
        check("Levels 1-3 do not",
              "Introduction" in paragraphs and "Prior Work" in paragraphs, True)
        check("the references page exists", "References" in paragraphs, True)

        indents = {p.text: (p.paragraph_format.first_line_indent,
                            p.paragraph_format.left_indent)
                   for p in document.paragraphs if p.text}
        check("body paragraphs are indented half an inch",
              indents["A paragraph of text."][0].inches, 0.5)
        quote = next(t for t in indents if t.startswith("A quoted passage"))
        check("the block quote is indented from the left, not the first line",
              (indents[quote][0], indents[quote][1].inches), (None, 0.5))
        entry = next(t for t in indents if t.startswith("Aiken, L."))
        check("the reference hangs",
              (indents[entry][0].inches, indents[entry][1].inches), (-0.5, 0.5))


def main() -> int:
    for name, function in sorted(globals().items()):
        if name.startswith("test_") and callable(function):
            function()
    print(f"tests/test_compose.py: {CHECKS} checks, {len(FAILURES)} failures")
    for failure in FAILURES:
        print(f"  ✗ {failure}")
    return 1 if FAILURES else 0


if __name__ == "__main__":
    raise SystemExit(main())
