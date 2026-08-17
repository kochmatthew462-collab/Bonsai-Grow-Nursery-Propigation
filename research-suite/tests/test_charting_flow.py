"""
Tests for the interlocks, the audit chain, storage round-tripping, the specialty
bundles, and the two exported documents.

The interlock tests are written as pairs: the note that must be refused, and the
note that must then be accepted once the missing element is supplied. A block that
cannot be cleared is a block that gets worked around, so both halves matter.

Everything runs offline against a temporary directory. Nothing here touches the
network, and there is no HTTP client involved — the route layer is exercised
through the same functions it calls, so a failure points at the logic rather than
at a serialisation layer.

Run: python3 tests/test_charting_flow.py
"""

from __future__ import annotations

import sys
import tempfile
from datetime import timedelta
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document  # noqa: E402

from app.charting import (  # noqa: E402
    disclosure, export, interlocks, narrative, specialty, store as store_module,
    systems,
)
from app.charting.models import (  # noqa: E402
    Encounter, Entry, EntryKind, Intervention, Medication, ProviderNotification,
    Setting, Severity, Vitals, iso, now_utc,
)

FAILURES: list[str] = []
CHECKS = 0


def check(label: str, got, want) -> None:
    global CHECKS
    CHECKS += 1
    if got != want:
        FAILURES.append(f"{label}\n     got: {got!r}\n    want: {want!r}")


def contains(label: str, haystack: str, needle: str) -> None:
    global CHECKS
    CHECKS += 1
    if needle.lower() not in (haystack or "").lower():
        FAILURES.append(f"{label}\n    {needle!r} not in {haystack!r}")


def absent(label: str, haystack: str, needle: str) -> None:
    global CHECKS
    CHECKS += 1
    if needle.lower() in (haystack or "").lower():
        FAILURES.append(f"{label}\n    {needle!r} unexpectedly in {haystack!r}")


NOW = now_utc()


def blank_encounter(setting: Setting = Setting.MED_SURG, **kwargs) -> Encounter:
    return Encounter(encounter_id="e1", local_label="Bed 12", setting=setting,
                     created_at=iso(NOW), updated_at=iso(NOW), **kwargs)


def note(**kwargs) -> Entry:
    entry = Entry(kind=kwargs.pop("kind", EntryKind.FOCUSED_UPDATE),
                  author_initials="MK", **kwargs)
    entry.stamp(NOW)
    return entry


def gate_codes(gate) -> set[str]:
    return {flag.code for flag in gate.flags}


def blocking_codes(gate) -> set[str]:
    return {f.code for f in gate.flags if f.severity is Severity.BLOCK}


# =================================================== critical value interlock


def test_critical_value_requires_notification() -> None:
    """The escalation matrix: the specification's grey-out requirement."""
    encounter = blank_encounter()
    entry = note()
    entry.vitals = Vitals(systolic=84, diastolic=50, heart_rate=134,
                          taken_at=iso(NOW))

    gate = interlocks.evaluate(encounter, entry, at=NOW)
    check("blocked", gate.blocked, True)
    check("notification required", gate.requires_notification, True)
    check("the block is the notification one",
          blocking_codes(gate), {"critical_value_unnotified"})
    check("three thresholds crossed", len(gate.critical_values), 3)
    contains("heart rate named", " ".join(gate.critical_values),
             "heart rate above 130")
    contains("calculated MAP counts", " ".join(gate.critical_values),
             "mean arterial pressure below 65")

    # Now supply the notification.
    entry.notifications.append(ProviderNotification(
        notified_at=iso(NOW), provider_name="Dr. Alvarez",
        provider_role="hospitalist", method="paged",
        reason="a systolic of 84 with a heart rate of 134",
        orders_received="500 mL lactated Ringer bolus", read_back=True,
        response="bolus then recheck in 30 minutes", response_at=iso(NOW)))
    gate = interlocks.evaluate(encounter, entry, at=NOW)
    check("cleared by a complete notification", gate.blocked, False)
    check("no blocks remain", blocking_codes(gate), set())


def test_escalation_counts_instead_of_notification() -> None:
    """A nurse who escalated rather than paged has still discharged the duty."""
    encounter = blank_encounter()
    entry = note()
    entry.vitals = Vitals(spo2=86, taken_at=iso(NOW))
    entry.notifications.append(ProviderNotification(
        notified_at=iso(NOW), provider_name="", method="",
        escalated_to="rapid response team", escalated_at=iso(NOW),
        escalation_response="Team at the bedside at 1420"))
    gate = interlocks.evaluate(encounter, entry, at=NOW)
    check("escalation satisfies the interlock", gate.blocked, False)


def test_stalled_chain_prompts_the_next_rung() -> None:
    encounter = blank_encounter()
    entry = note()
    entry.vitals = Vitals(systolic=84, taken_at=iso(NOW))
    entry.notifications.append(ProviderNotification(
        notified_at=iso(NOW), provider_name="Dr. Chen", provider_role="resident",
        method="paged", reason="a systolic of 84", no_response=True))
    gate = interlocks.evaluate(encounter, entry, at=NOW)
    check("not blocked — the nurse did notify", gate.blocked, False)
    check("but the chain prompt fires",
          "escalation_available" in gate_codes(gate), True)
    prompt = next(f for f in gate.flags if f.code == "escalation_available")
    contains("names a next rung", prompt.suggestion, "charge nurse")
    contains("says the chain does not end", prompt.message,
             "does not end here")

    # Once escalated, the prompt goes quiet.
    entry.notifications[-1].escalated_to = "charge nurse"
    entry.notifications[-1].escalated_at = iso(NOW)
    gate = interlocks.evaluate(encounter, entry, at=NOW)
    check("prompt clears after escalation",
          "escalation_available" in gate_codes(gate), False)


def test_thresholds_do_not_fire_on_normal_vitals() -> None:
    encounter = blank_encounter()
    entry = note()
    entry.vitals = Vitals(systolic=118, diastolic=72, heart_rate=78,
                          respiratory_rate=16, spo2=97, temperature_c=36.8,
                          pain_score=2, blood_glucose=104, taken_at=iso(NOW))
    gate = interlocks.evaluate(encounter, entry, at=NOW)
    check("nothing crossed", gate.critical_values, [])
    check("not blocked", gate.blocked, False)
    check("no notification required", gate.requires_notification, False)


# ================================================================ the clock


def test_future_event_is_refused_and_cannot_be_overridden() -> None:
    """Charting ahead of the clock is the shape of a pre-signed assessment."""
    encounter = blank_encounter()
    entry = Entry(kind=EntryKind.FOCUSED_UPDATE,
                  event_at=iso(NOW + timedelta(hours=1)))
    entry.stamp(NOW)
    gate = interlocks.evaluate(encounter, entry, at=NOW)
    check("blocked", gate.blocked, True)
    check("the right block", "event_in_future" in blocking_codes(gate), True)
    check("not overridable", gate.overridable, False)

    with_reason = interlocks.evaluate(encounter, entry, at=NOW,
                                      override_reason="I need to")
    check("an override does not clear it", with_reason.blocked, True)


def test_late_entry_is_labelled_not_refused() -> None:
    encounter = blank_encounter()
    entry = Entry(kind=EntryKind.EVENT_NOTE,
                  event_at=iso(NOW - timedelta(minutes=95)))
    entry.stamp(NOW)
    gate = interlocks.evaluate(encounter, entry, at=NOW)
    check("not blocked", gate.blocked, False)
    check("labelled", "late_entry" in gate_codes(gate), True)
    flag = next(f for f in gate.flags if f.code == "late_entry")
    check("informational only", flag.severity, Severity.INFO)
    contains("says nothing to fix", flag.suggestion, "Nothing to fix")

    # Just inside the window, no label.
    early = Entry(event_at=iso(NOW - timedelta(minutes=30)))
    early.stamp(NOW)
    check("30 minutes is not a late entry", early.late_entry, False)


def test_documentation_time_cannot_be_supplied() -> None:
    """`stamp()` is the only writer of `documented_at`, and it uses the clock."""
    entry = Entry(documented_at="1999-01-01T00:00:00+00:00")
    entry.stamp(NOW)
    check("the supplied value is overwritten", entry.documented_at, iso(NOW))


# ============================================================ incident report


def test_incident_report_block_is_absolute() -> None:
    encounter = blank_encounter()
    entry = note()
    entry.assessment = "Completed an incident report for the event."
    gate = interlocks.evaluate(encounter, entry, at=NOW,
                               override_reason="this is fine")
    check("blocked despite the override", gate.blocked, True)
    check("not overridable", gate.overridable, False)
    check("the right block", "incident_report" in blocking_codes(gate), True)


# ============================================================ copy-forward


def test_unchanged_copy_forward_is_refused() -> None:
    """A shift of byte-identical notes establishes that none of them describe an
    actual assessment."""
    encounter = blank_encounter()
    first = note(kind=EntryKind.SHIFT_ASSESSMENT)
    first.narrative = "Patient stable. No change from the previous assessment."
    first.vitals = Vitals(systolic=118, heart_rate=78, taken_at=iso(NOW))
    encounter.entries.append(first)

    duplicate = note(kind=EntryKind.SHIFT_ASSESSMENT)
    duplicate.narrative = first.narrative
    duplicate.vitals = Vitals(systolic=118, heart_rate=78, taken_at=iso(NOW))
    gate = interlocks.evaluate(encounter, duplicate, at=NOW)
    check("blocked", "copy_forward_unchanged" in blocking_codes(gate), True)
    check("overridable", gate.overridable, True)
    duplicate_flag = next(f for f in gate.flags
                          if f.code == "copy_forward_unchanged")
    contains("suggests the honest version", duplicate_flag.suggestion,
             "no change from the 1400 assessment")

    # Moving the vitals clears it: the note is now distinguishable.
    duplicate.vitals = Vitals(systolic=104, heart_rate=92, taken_at=iso(NOW))
    gate = interlocks.evaluate(encounter, duplicate, at=NOW)
    check("changed vitals clear it",
          "copy_forward_unchanged" in blocking_codes(gate), False)

    # So does changing the text.
    duplicate.vitals = Vitals(systolic=118, heart_rate=78, taken_at=iso(NOW))
    duplicate.narrative = first.narrative + " Ambulated 40 feet at 1500."
    gate = interlocks.evaluate(encounter, duplicate, at=NOW)
    check("changed text clears it",
          "copy_forward_unchanged" in blocking_codes(gate), False)

    # A different kind of note is never a copy-forward of this one.
    other_kind = note(kind=EntryKind.EVENT_NOTE)
    other_kind.narrative = first.narrative
    other_kind.vitals = Vitals(systolic=118, heart_rate=78, taken_at=iso(NOW))
    gate = interlocks.evaluate(encounter, other_kind, at=NOW)
    check("a different kind is not compared",
          "copy_forward_unchanged" in blocking_codes(gate), False)


# ======================================================== paediatric weight


def test_paediatric_weight_hard_stop() -> None:
    """A pounds-for-kilograms substitution is a 2.2-fold overdose."""
    encounter = blank_encounter(Setting.PEDIATRIC, age_band="4 y")
    entry = note()
    entry.medications.append(Medication(name="Paracetamol", dose="240 mg",
                                        route="oral", given_at=iso(NOW)))
    gate = interlocks.evaluate(encounter, entry, at=NOW)
    check("blocked", "paediatric_weight_required" in blocking_codes(gate), True)
    flag = next(f for f in gate.flags if f.code == "paediatric_weight_required")
    contains("insists on kilograms", flag.suggestion, "kilograms, not pounds")

    encounter.weight_kg = 16.0
    encounter.weight_measured_at = iso(NOW)
    gate = interlocks.evaluate(encounter, entry, at=NOW)
    check("cleared by a weight",
          "paediatric_weight_required" in blocking_codes(gate), False)

    # An adult encounter with no weight is not blocked.
    adult = blank_encounter()
    adult_entry = note()
    adult_entry.medications.append(Medication(name="Paracetamol", dose="1 g",
                                              given_at=iso(NOW)))
    adult_gate = interlocks.evaluate(adult, adult_entry, at=NOW)
    check("adults are not weight-gated",
          "paediatric_weight_required" in blocking_codes(adult_gate), False)

    # A note with no medication is not gated either, even in paediatrics.
    no_meds = note()
    check("no medication, no gate",
          "paediatric_weight_required" in blocking_codes(
              interlocks.evaluate(blank_encounter(Setting.PEDIATRIC), no_meds,
                                  at=NOW)),
          False)


def test_age_band_triggers_the_weight_rule() -> None:
    """A paediatric patient on an adult unit still needs the weight."""
    for band in ("6 months", "neonate", "preterm", "child", "8 y"):
        encounter = blank_encounter(Setting.MED_SURG, age_band=band)
        check(f"{band!r} requires a weight", encounter.requires_weight(), True)
    check("adult does not",
          blank_encounter(Setting.MED_SURG, age_band="adult").requires_weight(),
          False)


def test_stale_weight_warns() -> None:
    encounter = blank_encounter(Setting.PEDIATRIC, age_band="2 y")
    encounter.weight_kg = 12.0
    encounter.weight_measured_at = iso(NOW - timedelta(hours=40))
    entry = note()
    entry.medications.append(Medication(name="Ibuprofen", given_at=iso(NOW)))
    gate = interlocks.evaluate(encounter, entry, at=NOW)
    check("warns", "weight_stale" in gate_codes(gate), True)
    check("does not block", gate.blocked, False)


def test_high_alert_double_check() -> None:
    encounter = blank_encounter()
    entry = note()
    entry.medications.append(Medication(name="Insulin", dose="8 units",
                                        route="subcutaneous", given_at=iso(NOW),
                                        high_alert=True))
    gate = interlocks.evaluate(encounter, entry, at=NOW)
    check("warns", "high_alert_no_double_check" in gate_codes(gate), True)
    entry.medications[0].second_nurse_verified = True
    entry.medications[0].second_nurse_initials = "JP"
    gate = interlocks.evaluate(encounter, entry, at=NOW)
    check("cleared", "high_alert_no_double_check" in gate_codes(gate), False)


def test_controlled_substance_waste_needs_a_witness() -> None:
    encounter = blank_encounter()
    entry = note()
    entry.medications.append(Medication(name="Hydromorphone", dose="0.5 mg",
                                        given_at=iso(NOW), wasted_amount="0.5 mg"))
    gate = interlocks.evaluate(encounter, entry, at=NOW)
    check("warns", "waste_without_witness" in gate_codes(gate), True)
    flag = next(f for f in gate.flags if f.code == "waste_without_witness")
    contains("names the real consequence", flag.suggestion,
             "diversion investigation")


# ============================================================== closed loops


def test_analgesic_opens_a_thirty_minute_loop() -> None:
    """The specification's named example."""
    encounter = blank_encounter()
    entry = note()
    entry.medications.append(Medication(
        name="Morphine", dose="2 mg", route="IV", given_at=iso(NOW),
        indication="pain rated 7 of 10"))
    loops = interlocks.loops_for_entry(entry, at=NOW)
    check("one loop", len(loops), 1)
    check("thirty-minute window", loops[0].window_minutes, 30)
    contains("owed for pain reassessment", loops[0].trigger, "pain reassessment")


def test_loop_windows_by_class() -> None:
    encounter = blank_encounter()
    cases = [
        ("Norepinephrine", "maintaining a MAP above 65", 15),
        ("Naloxone", "opioid reversal", 15),
        ("Dextrose 50%", "hypoglycaemia", 15),
        ("Ondansetron", "nausea", 60),
        ("Insulin", "hyperglycaemia", 60),
        ("Albuterol", "wheeze", 30),
        ("Lorazepam", "agitation", 30),
    ]
    for name, indication, window in cases:
        entry = note()
        entry.medications.append(Medication(name=name, given_at=iso(NOW),
                                            indication=indication))
        loops = interlocks.loops_for_entry(entry, at=NOW)
        check(f"{name} opens a loop", len(loops) >= 1, True)
        check(f"{name} window is {window} minutes",
              min(l.window_minutes for l in loops), window)


def test_held_medication_opens_no_loop() -> None:
    entry = note()
    entry.medications.append(Medication(name="Morphine", held=True,
                                        hold_reason="respiratory rate 8",
                                        indication="pain"))
    check("no loop for a held dose", interlocks.loops_for_entry(entry, at=NOW), [])


def test_overdue_loop_reported() -> None:
    encounter = blank_encounter()
    # The window runs from the *event*, not from when the note was registered —
    # a dose given at 1400 owes its reassessment at 1430 however late the note was
    # written. So the event time is what has to be pushed back here.
    entry = note(event_at=iso(NOW - timedelta(minutes=90)))
    entry.medications.append(Medication(
        name="Morphine", given_at=iso(NOW - timedelta(minutes=90)),
        indication="pain"))
    interlocks.register_loops(encounter, entry, at=NOW)
    check("one loop open", len(encounter.open_loops()), 1)
    check("and overdue", len(encounter.overdue_loops(NOW)), 1)
    loop = encounter.reassessments[0]
    check("60 minutes late", loop.minutes_late(NOW), 60)

    flags = interlocks.overdue_summary(encounter, at=NOW)
    check("reported", len(flags), 1)
    contains("says how late", flags[0].message, "60 minutes past")
    contains("offers the late-entry route", flags[0].suggestion, "late entry")

    # Closing it with a finding clears it.
    interlocks.satisfy_loop(encounter, loop.reassessment_id,
                            finding="Reports 3 of 10, resting comfortably", at=NOW)
    check("satisfied", encounter.open_loops(), [])
    check("no longer overdue", encounter.overdue_loops(NOW), [])

    # Or with a reason it was not done.
    second = note(event_at=iso(NOW - timedelta(minutes=90)))
    second.medications.append(Medication(
        name="Hydromorphone", given_at=iso(NOW - timedelta(minutes=90)),
        indication="pain"))
    interlocks.register_loops(encounter, second, at=NOW)
    open_loop = encounter.open_loops()[0]
    interlocks.satisfy_loop(encounter, open_loop.reassessment_id, finding="",
                            not_done_reason="patient off the unit for imaging")
    check("not-done also closes it", encounter.open_loops(), [])


def test_loops_are_not_duplicated() -> None:
    encounter = blank_encounter()
    entry = note()
    entry.medications.append(Medication(name="Morphine", given_at=iso(NOW),
                                        indication="pain"))
    interlocks.register_loops(encounter, entry, at=NOW)
    interlocks.register_loops(encounter, entry, at=NOW)
    check("registered once", len(encounter.reassessments), 1)


# ========================================================== completeness warnings


def test_refusal_four_pillars() -> None:
    encounter = blank_encounter()
    entry = note(kind=EntryKind.REFUSAL)
    gate = interlocks.evaluate(encounter, entry, at=NOW)
    for pillar in ("capacity", "risks_explained", "provider_notified",
                   "patient_quote"):
        check(f"{pillar} reported missing",
              f"refusal_missing_{pillar}" in gate_codes(gate), True)

    entry.module_data["refusal"] = {
        "capacity": "oriented ×4, able to state the consequence",
        "risks_explained": "rebound hypertension",
        "provider_notified": "Dr. Osei at 0915",
        "patient_quote": "I am not taking that",
    }
    gate = interlocks.evaluate(encounter, entry, at=NOW)
    check("all four cleared",
          [c for c in gate_codes(gate) if c.startswith("refusal_missing")], [])


def test_restraint_requirements() -> None:
    encounter = blank_encounter()
    entry = note(kind=EntryKind.RESTRAINT)
    gate = interlocks.evaluate(encounter, entry, at=NOW)
    for key in ("restraint_no_order", "restraint_missing_checks",
                "restraint_missing_circulation", "restraint_missing_release_rom",
                "restraint_missing_needs_offered",
                "restraint_missing_discontinuation_criteria"):
        check(f"{key} fires", key in gate_codes(gate), True)


def test_blood_product_requirements() -> None:
    encounter = blank_encounter()
    entry = note(kind=EntryKind.BLOOD_PRODUCT)
    gate = interlocks.evaluate(encounter, entry, at=NOW)
    for key in ("blood_missing_consent", "blood_missing_two_person_verification",
                "blood_missing_baseline_vitals", "blood_missing_unit_number"):
        check(f"{key} fires", key in gate_codes(gate), True)
    flag = next(f for f in gate.flags
                if f.code == "blood_missing_two_person_verification")
    contains("names why it matters", flag.suggestion, "prevents the error that kills")


def test_fall_requirements() -> None:
    encounter = blank_encounter()
    entry = note(kind=EntryKind.FALL_EVENT)
    gate = interlocks.evaluate(encounter, entry, at=NOW)
    check("witnessed status required",
          "fall_witnessed_unknown" in gate_codes(gate), True)
    check("neuro checks required",
          "fall_no_neuro_checks" in gate_codes(gate), True)
    flag = next(f for f in gate.flags if f.code == "fall_no_neuro_checks")
    contains("names anticoagulation", flag.suggestion, "anticoagulated")


def test_notification_outcome_and_read_back() -> None:
    encounter = blank_encounter()
    entry = note()
    entry.notifications.append(ProviderNotification(
        notified_at=iso(NOW), provider_name="Dr. Alvarez", method="paged",
        reason="a change in status"))
    gate = interlocks.evaluate(encounter, entry, at=NOW)
    check("no outcome warns", "notification_no_outcome" in gate_codes(gate), True)

    entry.notifications[0].orders_received = "increase oxygen to 4 L/min"
    gate = interlocks.evaluate(encounter, entry, at=NOW)
    check("outcome clears it", "notification_no_outcome" in gate_codes(gate), False)
    check("but read-back is now asked for",
          "orders_not_read_back" in gate_codes(gate), True)
    entry.notifications[0].read_back = True
    gate = interlocks.evaluate(encounter, entry, at=NOW)
    check("read-back clears it", "orders_not_read_back" in gate_codes(gate), False)


def test_intervention_needs_a_response() -> None:
    encounter = blank_encounter()
    entry = note()
    entry.interventions.append(Intervention(action="Repositioned to the left side",
                                            performed_at=iso(NOW)))
    gate = interlocks.evaluate(encounter, entry, at=NOW)
    check("warns", "intervention_not_evaluated" in gate_codes(gate), True)
    entry.interventions[0].response = "Tolerated, no change in saturation"
    gate = interlocks.evaluate(encounter, entry, at=NOW)
    check("cleared", "intervention_not_evaluated" in gate_codes(gate), False)


def test_shift_assessment_unassessed_systems() -> None:
    encounter = blank_encounter()
    entry = note(kind=EntryKind.SHIFT_ASSESSMENT)
    gate = interlocks.evaluate(encounter, entry, at=NOW)
    check("all eight reported", "systems_unassessed" in gate_codes(gate), True)
    flag = next(f for f in gate.flags if f.code == "systems_unassessed")
    contains("counts them", flag.message, "8 body systems")

    entry.systems = [systems.build_finding(s.system_id, within_defined_limits=True)
                     for s in systems.SYSTEMS]
    gate = interlocks.evaluate(encounter, entry, at=NOW)
    check("cleared", "systems_unassessed" in gate_codes(gate), False)
    check("but rubber-stamping is flagged",
          "wdl_across_the_board" in gate_codes(gate), True)
    rubber = next(f for f in gate.flags if f.code == "wdl_across_the_board")
    contains("acknowledges it may be accurate", rubber.suggestion,
             "may be entirely accurate")


# ================================================================ the override


def test_override_records_itself() -> None:
    encounter = blank_encounter()
    entry = note()
    entry.vitals = Vitals(spo2=86, taken_at=iso(NOW))

    without = interlocks.evaluate(encounter, entry, at=NOW)
    check("blocked without a reason", without.blocked, True)
    check("overridable", without.overridable, True)

    with_reason = interlocks.evaluate(
        encounter, entry, at=NOW,
        override_reason="Rapid response in progress; completing after")
    check("saves with a reason", with_reason.blocked, False)
    check("and records it", "override_recorded" in gate_codes(with_reason), True)
    recorded = next(f for f in with_reason.flags if f.code == "override_recorded")
    contains("carries the reason", recorded.excerpt, "Rapid response in progress")
    contains("asks for completion", recorded.suggestion, "addendum")


# =========================================================== non-destructive edit


def test_revision_keeps_the_original() -> None:
    entry = note()
    entry.narrative = "Original text."
    revision = entry.revise("Corrected text.", "typographical error", "MK")
    check("narrative updated", entry.narrative, "Corrected text.")
    check("original kept", revision.previous_text, "Original text.")
    check("reason kept", revision.reason, "typographical error")
    check("one revision", len(entry.revisions), 1)

    entry.revise("Corrected again.", "clarification", "MK")
    check("appended not replaced", len(entry.revisions), 2)
    check("the chain is walkable",
          [r.previous_text for r in entry.revisions],
          ["Original text.", "Corrected text."])


def test_revision_requires_a_reason() -> None:
    entry = note()
    entry.narrative = "Text."
    for reason in ("", "   ", "\n"):
        raised = False
        try:
            entry.revise("New text.", reason)
        except ValueError as error:
            raised = True
            contains("the error explains why", str(error),
                     "indefensible")
        check(f"a reason of {reason!r} is refused", raised, True)
    check("nothing was changed", entry.narrative, "Text.")
    check("no revision recorded", entry.revisions, [])


# ================================================================= audit chain


def test_audit_chain_detects_alteration() -> None:
    encounter = blank_encounter()
    encounter.record("encounter_created", "e1", "", "MK")
    encounter.record("entry_created", "x1", "shift assessment", "MK")
    encounter.record("entry_revised", "x1", "typographical error", "MK")
    result = store_module.verify_chain(encounter)
    check("intact", result["intact"], True)
    check("three events verified", result["verified"], 3)
    contains("honest about what it proves", result["reason"],
             "not tamper-proofing")

    # Altering an event's content breaks its own digest.
    encounter.audit[1].detail = "something else"
    broken = store_module.verify_chain(encounter)
    check("alteration detected", broken["intact"], False)
    check("at the right index", broken["broken_at"], 1)
    contains("says it was altered", broken["reason"], "altered after it was written")


def test_audit_chain_detects_deletion() -> None:
    encounter = blank_encounter()
    for index in range(4):
        encounter.record("event", f"t{index}", "", "MK")
    del encounter.audit[2]
    result = store_module.verify_chain(encounter)
    check("deletion detected", result["intact"], False)
    check("at the gap", result["broken_at"], 2)
    contains("says it was removed", result["reason"], "inserted or removed")


def test_empty_chain_verifies() -> None:
    check("an encounter with no events is intact",
          store_module.verify_chain(blank_encounter())["intact"], True)


# =================================================================== storage


def test_storage_round_trip() -> None:
    directory = Path(tempfile.mkdtemp())
    store = store_module.EncounterStore(directory / "charting")
    encounter = store.new_encounter("Bed 12", setting="icu", initials="AB",
                                    author_initials="MK")
    encounter.weight_kg = 78.5
    encounter.code_status = "Full code"
    encounter.context.patient_count = 6
    encounter.lines_drains_airways = [
        {"kind": "Central venous catheter", "site": "right internal jugular",
         "inserted_at": "2026-08-15 14:00", "indication": "vasopressor infusion"}]

    entry = note(kind=EntryKind.SHIFT_ASSESSMENT)
    entry.vitals = Vitals(systolic=118, diastolic=72, heart_rate=84,
                          taken_at=iso(NOW))
    entry.systems = [systems.build_finding("neuro", within_defined_limits=True)]
    entry.medications = [Medication(name="Morphine", dose="2 mg",
                                    given_at=iso(NOW), indication="pain")]
    entry.notifications = [ProviderNotification(
        notified_at=iso(NOW), provider_name="Dr. Alvarez", method="paged",
        reason="pain control")]
    entry.narrative = "Composed note."
    entry.revise("Corrected note.", "typographical error", "MK")
    encounter.entries.append(entry)
    interlocks.register_loops(encounter, entry, at=NOW)
    store.save(encounter)

    loaded = store.load(encounter.encounter_id)
    check("setting survives", loaded.setting, Setting.ICU)
    check("weight survives", loaded.weight_kg, 78.5)
    check("code status survives", loaded.code_status, "Full code")
    check("shift context survives", loaded.context.patient_count, 6)
    check("devices survive", len(loaded.lines_drains_airways), 1)
    check("entry kind survives", loaded.entries[0].kind,
          EntryKind.SHIFT_ASSESSMENT)
    check("vitals survive", loaded.entries[0].vitals.heart_rate, 84.0)
    check("systems survive", loaded.entries[0].systems[0].system, "neuro")
    check("wdl flag survives",
          loaded.entries[0].systems[0].within_defined_limits, True)
    check("medications survive", loaded.entries[0].medications[0].name, "Morphine")
    check("notifications survive",
          loaded.entries[0].notifications[0].provider_name, "Dr. Alvarez")
    check("revisions survive", len(loaded.entries[0].revisions), 1)
    check("superseded text survives",
          loaded.entries[0].revisions[0].previous_text, "Composed note.")
    check("loops survive", len(loaded.reassessments), 1)
    check("chain still verifies",
          store_module.verify_chain(loaded)["intact"], True)

    listed = store.list_encounters()
    check("listed once", len(listed), 1)
    check("untranscribed counted", listed[0]["untranscribed"], 1)

    check("purge removes it", store.purge_all(), 1)
    check("and the directory is empty", store.list_encounters(), [])


def test_unknown_keys_are_ignored_on_load() -> None:
    """A file written by a later version must still load."""
    directory = Path(tempfile.mkdtemp())
    store = store_module.EncounterStore(directory / "charting")
    encounter = store.new_encounter("Bed 3", author_initials="MK")
    path = store.path_for(encounter.encounter_id)
    import json
    raw = json.loads(path.read_text("utf-8"))
    raw["some_future_field"] = {"a": 1}
    raw["entries"] = [{"entry_id": "e", "kind": "focused_update",
                       "narrative": "x", "unknown_thing": True}]
    path.write_text(json.dumps(raw), "utf-8")
    loaded = store.load(encounter.encounter_id)
    check("loaded despite unknown keys", loaded.local_label, "Bed 3")
    check("entry loaded", loaded.entries[0].narrative, "x")


def test_unsafe_ids_are_refused() -> None:
    directory = Path(tempfile.mkdtemp())
    store = store_module.EncounterStore(directory / "charting")
    for bad in ("..", ".", "../../etc/passwd", "/", ""):
        raised = False
        try:
            store.path_for(bad)
        except ValueError:
            raised = True
        check(f"{bad!r} refused", raised, True)


# ================================================================== bundles


def test_sepsis_bundle_intervals() -> None:
    status = specialty.bundle_status("sepsis", {
        "time_zero": "1400", "lactate": "1425", "blood_cultures": "1430",
        "antibiotics": "1745", "fluids": "1440"})
    check("anchor recorded", status["anchor_recorded"], True)
    check("one target exceeded", len(status["exceeded"]), 1)
    contains("the antibiotic is the one", status["exceeded"][0], "antibiotics")
    check("three elements outstanding", len(status["outstanding"]), 3)
    antibiotic = next(r for r in status["rows"] if r["key"] == "antibiotics")
    check("225 minutes", antibiotic["interval_minutes"], 225)


def test_bundle_without_an_anchor_computes_nothing() -> None:
    """This is the whole point of the sepsis bundle's own note."""
    status = specialty.bundle_status("sepsis", {"lactate": "1425",
                                               "antibiotics": "1500"})
    check("no anchor", status["anchor_recorded"], False)
    check("no intervals",
          [r["interval_minutes"] for r in status["rows"] if r["interval_minutes"]],
          [])
    contains("says so", status["note"], "no interval in this bundle can be computed")


def test_stroke_anchors_on_arrival_not_last_known_well() -> None:
    """Door-to-imaging runs from arrival. Measuring it from last known well would
    report a compliant department as non-compliant."""
    status = specialty.bundle_status("stroke", {
        "last_known_well": "2330", "arrival": "0015", "stroke_alert": "0020",
        "ct_started": "0040"})
    check("anchored on arrival", status["anchor_label"], "Arrival time")
    ct = next(r for r in status["rows"] if r["key"] == "ct_started")
    check("25 minutes from arrival, not 70 from last known well",
          ct["interval_minutes"], 25)
    check("within the 25-minute target", ct["exceeded"], False)


def test_bundle_handles_midnight() -> None:
    status = specialty.bundle_status("stemi", {"arrival": "2350",
                                              "first_ecg": "2358"})
    ecg = next(r for r in status["rows"] if r["key"] == "first_ecg")
    check("eight minutes across midnight-adjacent times",
          ecg["interval_minutes"], 8)

    across = specialty.bundle_status("stemi", {"arrival": "2350",
                                              "first_ecg": "0005"})
    ecg = next(r for r in across["rows"] if r["key"] == "first_ecg")
    check("fifteen minutes across midnight", ecg["interval_minutes"], 15)


def test_bundle_time_parsing() -> None:
    check("HHMM", specialty._minutes("1412"), 14 * 60 + 12)
    check("HH:MM", specialty._minutes("14:12"), 14 * 60 + 12)
    check("HMM", specialty._minutes("912"), 9 * 60 + 12)
    check("blank", specialty._minutes(""), None)
    check("nonsense", specialty._minutes("later"), None)
    check("impossible hour", specialty._minutes("2512"), None)
    check("impossible minute", specialty._minutes("1499"), None)


def test_bundle_renders_missing_elements() -> None:
    text = specialty.render_bundle("sepsis", {"time_zero": "1400"})
    contains("names what is not recorded", text, "not recorded")
    contains("names the bundle", text, "Sepsis bundle")


# ================================================================== modules


def test_modules_cover_the_four_settings() -> None:
    for setting in (Setting.EMERGENCY, Setting.ICU, Setting.PEDIATRIC,
                    Setting.MED_SURG):
        module = specialty.module_for(setting)
        check(f"{setting.value} has a module", module is not None, True)
        check(f"{setting.value} module has blocks", len(module.blocks) >= 3, True)
        check(f"{setting.value} names priority scales",
              len(module.priority_scales) >= 4, True)

    check("an unknown setting has none", specialty.module_for("nonsense"), None)


def test_icu_titration_requires_the_reason() -> None:
    """The block that exists because flowsheets capture the rate and lose the why."""
    missing = specialty.missing_required(Setting.ICU, "titration", {})
    for label in ("Parameter titrated to", "Reading that prompted the change",
                  "Response after the change"):
        check(f"{label} required", label in missing, True)

    filled = specialty.missing_required(Setting.ICU, "titration", {
        "time": "1420", "drug": "Norepinephrine", "previous_rate": "6 mcg/min",
        "new_rate": "8 mcg/min", "direction": "increased",
        "parameter": "MAP above 65", "reading": "MAP 58",
        "response": "MAP 71 at 1435"})
    check("nothing missing when filled", filled, [])


def test_emergency_requires_time_at_the_bedside() -> None:
    module = specialty.module_for(Setting.EMERGENCY)
    timing = next(b for b in module.blocks if b.block_id == "timing")
    bedside = next(f for f in timing.fields
                   if f.key == "first_provider_contact")
    contains("distinguishes it from the page time", bedside.hint,
             "not the time they were paged")


# ================================================================== handoff


def test_handoff_names_what_was_not_handed_over() -> None:
    encounter = blank_encounter(Setting.EMERGENCY, code_status="Full code",
                                allergies="Penicillin — rash")
    encounter.lines_drains_airways = [
        {"kind": "Peripheral intravenous catheter", "site": "left antecubital",
         "size": "18 gauge", "inserted_at": "1350"}]
    entry = note()
    entry.medications.append(Medication(name="Morphine", given_at=iso(NOW),
                                        indication="pain"))
    interlocks.register_loops(encounter, entry, at=NOW)
    encounter.entries.append(entry)

    text = narrative.render_handoff("er_to_icu", {
        "code_status": "Full code, confirmed with the patient",
        "airway": "Patent, maintaining own airway"}, encounter)

    contains("carries the label", text, "Bed 12")
    contains("carries allergies", text, "Penicillin")
    contains("filled items present", text, "Patent, maintaining own airway")
    contains("unfilled required items are named", text, "NOT HANDED OVER")
    contains("the omission list is summarised", text,
             "Required items not handed over")
    contains("devices are listed", text, "18 gauge")
    contains("owed reassessments are listed", text, "Reassessments owed")
    contains("untranscribed notes are called out", text, "not yet transcribed")

    routes = {r["route"] for r in narrative.handoff_routes()}
    check("four routes", len(routes), 4)
    for route in routes:
        template = narrative.handoff_template(route)
        check(f"{route} has fields", len(template["fields"]) >= 5, True)
        check(f"{route} explains itself", bool(template["why"].strip()), True)


# ==================================================================== export


def test_exported_documents() -> None:
    directory = Path(tempfile.mkdtemp())
    store = store_module.EncounterStore(directory / "charting")
    encounter = store.new_encounter("Bed 12", setting="icu", author_initials="MK")
    encounter.code_status = "Full code"

    transcribed = note(kind=EntryKind.FOCUSED_UPDATE)
    transcribed.narrative = "Already filed in the chart."
    transcribed.transcribed_at = iso(NOW)
    outstanding = note(kind=EntryKind.EVENT_NOTE)
    outstanding.narrative = "Not yet filed."
    outstanding.revise("Not yet filed, corrected.", "typographical error", "MK")
    outstanding.flags = [{"code": "late_entry", "severity": "info",
                          "message": "tagged late", "field_path": "event_at",
                          "excerpt": "", "suggestion": ""}]
    encounter.entries += [transcribed, outstanding]
    encounter.record("entry_created", outstanding.entry_id, "event note", "MK")
    store.save(encounter)

    produced = export.write_documents(encounter, directory / "exports",
                                      author="MK, RN")
    check("two documents", len(produced), 2)

    record = Document(str(directory / "exports" / produced[0]["name"]))
    paragraphs = [p.text for p in record.paragraphs if p.text.strip()]
    joined = "\n".join(paragraphs)
    contains("carries the not-the-record header", joined,
             "not the legal medical record")
    contains("untranscribed section present", joined, "Not yet transcribed")
    contains("transcribed section present", joined, "Transcribed notes")
    check("untranscribed comes first",
          joined.index("Not yet transcribed") < joined.index("Transcribed notes"),
          True)

    struck = [r.text for p in record.paragraphs for r in p.runs if r.font.strike]
    check("superseded text is struck through", struck, ["Not yet filed."])

    audit_doc = Document(str(directory / "exports" / produced[1]["name"]))
    audit_text = "\n".join(p.text for p in audit_doc.paragraphs if p.text.strip())
    for heading in ("1. Timestamp ledger", "2. Corrections",
                    "3. Interlocks, warnings and overrides",
                    "4. Scoring instruments used",
                    "5. Audit chain verification",
                    "6. What this record does not establish"):
        contains(f"{heading} present", audit_text, heading)
    contains("the ledger explains the timestamp pair", audit_text,
             "cannot be set by any user")
    audit_struck = [r.text for p in audit_doc.paragraphs
                    for r in p.runs if r.font.strike]
    check("the appendix also strikes the superseded text",
          audit_struck, ["Not yet filed."])
    check("appendix has the three tables", len(audit_doc.tables), 3)


# ================================================================ disclosure


def test_disclosure_is_stated_everywhere_it_matters() -> None:
    banner = disclosure.banner()
    for key in ("not_the_record", "no_phi", "not_clinical_advice", "local_only"):
        check(f"{key} present", bool(banner[key].strip()), True)
    contains("names the EHR as the record", banner["not_the_record"], "EHR")
    contains("tells the nurse to transcribe", banner["not_the_record"],
             "transcribe")
    contains("no PHI is stated as a design fact", banner["no_phi"],
             "no fields for them")

    ledger = disclosure.what_it_cannot_do()
    check("six honest entries", len(ledger), 6)
    guarantee = ledger[0]
    contains("the first is the coverage claim", guarantee["claim"], "covered")
    contains("and it says nothing can", guarantee["reality"], "Nothing can")

    contains("the export header names the governing record",
             disclosure.EXPORT_HEADER, "the EHR governs")



# ================================================ regressions found by audit


def test_paediatric_weight_is_a_hard_stop() -> None:
    """The specification called it a hard stop and it was overridable: any typed
    reason cleared it. There is no emergency in which the right answer is to
    document a paediatric dose without the weight it was calculated from."""
    encounter = blank_encounter(Setting.PEDIATRIC, age_band="4 y")
    entry = note()
    entry.medications.append(Medication(name="Paracetamol", given_at=iso(NOW)))
    gate = interlocks.evaluate(encounter, entry, at=NOW,
                               override_reason="emergency, will add later")
    check("still blocked", gate.blocked, True)
    check("and not overridable", gate.overridable, False)
    check("listed as non-overridable",
          "paediatric_weight_required" in interlocks.NON_OVERRIDABLE, True)


def test_notification_without_an_outcome_does_not_clear_the_gate() -> None:
    """The specification asked for time, name, method AND response. Checking only
    the first three let the "MD aware" note the interlock exists to prevent
    straight through."""
    encounter = blank_encounter()
    entry = note()
    entry.vitals = Vitals(systolic=84, taken_at=iso(NOW))
    entry.notifications.append(ProviderNotification(
        notified_at=iso(NOW), provider_name="Dr. Alvarez", method="paged",
        reason="a systolic of 84"))
    check("still blocked", interlocks.evaluate(encounter, entry, at=NOW).blocked,
          True)

    entry.notifications[0].orders_received = "500 mL bolus"
    check("orders clear it",
          interlocks.evaluate(encounter, entry, at=NOW).blocked, False)

    entry.notifications[0].orders_received = ""
    entry.notifications[0].no_new_orders = True
    check("an explicit no-orders also clears it",
          interlocks.evaluate(encounter, entry, at=NOW).blocked, False)


def test_copy_forward_is_not_cleared_by_a_timestamp() -> None:
    """Comparing the rendered vitals sentence meant bumping only "time taken"
    cleared the duplicate check while every reading stayed identical."""
    encounter = blank_encounter()
    first = note(kind=EntryKind.SHIFT_ASSESSMENT)
    first.narrative = "Patient stable."
    first.vitals = Vitals(systolic=118, heart_rate=78, taken_at=iso(NOW))
    encounter.entries.append(first)

    same = note(kind=EntryKind.SHIFT_ASSESSMENT)
    same.narrative = "Patient stable."
    same.vitals = Vitals(systolic=118, heart_rate=78,
                         taken_at=iso(NOW + timedelta(hours=4)))
    check("a new timestamp alone does not clear it",
          "copy_forward_unchanged" in blocking_codes(
              interlocks.evaluate(encounter, same, at=NOW)), True)

    same.vitals.heart_rate = 92
    check("a changed reading does",
          "copy_forward_unchanged" in blocking_codes(
              interlocks.evaluate(encounter, same, at=NOW)), False)


def test_no_orders_prompts_escalation_without_a_checkbox() -> None:
    """Conditioning the prompt on the nurse ticking `no_new_orders` meant the
    commonest real case — a provider who answered and ordered nothing, on a
    patient still crossing a threshold — never prompted."""
    encounter = blank_encounter()
    entry = note()
    entry.vitals = Vitals(systolic=84, taken_at=iso(NOW))
    entry.notifications.append(ProviderNotification(
        notified_at=iso(NOW), provider_name="Dr. Chen", method="paged",
        reason="a systolic of 84", response="will review shortly"))
    check("prompted", "escalation_available" in gate_codes(
        interlocks.evaluate(encounter, entry, at=NOW)), True)


def test_reevaluation_timed_to_the_dose_is_flagged() -> None:
    encounter = blank_encounter()
    entry = note()
    entry.medications.append(Medication(
        name="Morphine", given_at=iso(NOW), indication="pain",
        effect="reports 3 of 10",
        effect_checked_at=iso(NOW + timedelta(minutes=1))))
    check("flagged", "effect_checked_too_soon" in gate_codes(
        interlocks.evaluate(encounter, entry, at=NOW)), True)

    entry.medications[0].effect_checked_at = iso(NOW + timedelta(minutes=35))
    check("a real re-evaluation is not",
          "effect_checked_too_soon" in gate_codes(
              interlocks.evaluate(encounter, entry, at=NOW)), False)


def test_procedure_notes_have_completeness_rules() -> None:
    """'Procedure' was only an entry-kind label with no block behind it."""
    encounter = blank_encounter()
    entry = note(kind=EntryKind.PROCEDURE)
    codes = gate_codes(interlocks.evaluate(encounter, entry, at=NOW))
    check("an empty procedure note is flagged",
          "procedure_note_without_procedure" in codes, True)

    entry.module_data["procedures"] = [{"name": "Peripheral IV insertion",
                                        "performed_at": iso(NOW)}]
    codes = gate_codes(interlocks.evaluate(encounter, entry, at=NOW))
    for key in ("procedure_missing_consent", "procedure_missing_tolerance",
                "procedure_missing_post_assessment"):
        check(f"{key} fires", key in codes, True)

    entry.module_data["procedures"][0].update(
        consent="Verbal consent obtained",
        tolerance="Tolerated well, no complaint of pain",
        post_assessment="Site clean and dry, flushes freely, no swelling")
    codes = gate_codes(interlocks.evaluate(encounter, entry, at=NOW))
    check("all cleared",
          [c for c in codes if c.startswith("procedure_missing")], [])


def test_reason_for_the_note_reaches_the_subjective_section() -> None:
    """The specification put "the reason for this update" in S beside the
    patient's words, and there was no field for it."""
    encounter = blank_encounter()
    entry = note()
    entry.reason = "Called to the room for new chest tightness"
    entry.subjective = 'Patient stated, "It feels tight."'
    text = narrative.compose(entry, encounter)
    contains("reason present", text, "Reason for this note")
    contains("and the patient's words", text, "It feels tight")
    check("reason precedes the quote",
          text.index("Reason for this note") < text.index("It feels tight"), True)


def main() -> int:
    for name, function in sorted(globals().items()):
        if name.startswith("test_") and callable(function):
            function()
    print(f"tests/test_charting_flow.py: {CHECKS} checks, {len(FAILURES)} failures")
    for failure in FAILURES:
        print(f"  ✗ {failure}")
    return 1 if FAILURES else 0


if __name__ == "__main__":
    raise SystemExit(main())
