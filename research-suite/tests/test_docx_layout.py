"""
APA 7 layout tests against the generated OOXML.

These assert on the XML of a real saved document rather than on the builder's
own return values, because everything that goes wrong in Word export goes wrong
between the API call and the file: a font set on one of four attributes, a page
number written as a literal digit, a hanging indent applied to the wrong
paragraph. A test that only checks the builder did not raise would pass on all
three.

Run: python3 tests/test_docx_layout.py
"""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Inches  # noqa: E402

from app.apa.citations import CitationContext, Run  # noqa: E402
from app.apa.document import ApaPaper, Block, default_running_head  # noqa: E402
from app.models import Author, Project, TitlePage, Work, WorkType  # noqa: E402

FAILURES: list[str] = []
CHECKS = 0
EMU_PER_INCH = 914400


def check(label: str, got, want) -> None:
    global CHECKS
    CHECKS += 1
    if got != want:
        FAILURES.append(f"{label}\n     got: {got!r}\n    want: {want!r}")


def near(label: str, got, want, tol=1000) -> None:
    global CHECKS
    CHECKS += 1
    if got is None or abs(int(got) - int(want)) > tol:
        FAILURES.append(f"{label}\n     got: {got!r}\n    want: ~{want!r}")


def paragraphs_of(path) -> list:
    """Snapshot the paragraph list once.

    `Document.paragraphs` constructs new wrapper objects on every access, so
    `doc.paragraphs.index(p)` compares a fresh object against a stale one and
    always raises. Everything below indexes into one captured list instead.
    """
    return list(Document(str(path)).paragraphs)


def find_index(paras: list, predicate) -> int:
    for index, paragraph in enumerate(paras):
        if predicate(paragraph):
            return index
    raise AssertionError("paragraph not found")


def sample_project(variant: str = "student") -> Project:
    work = Work(
        work_type=WorkType.JOURNAL_ARTICLE,
        authors=[Author.parse("Smith, John A."), Author.parse("Jones, Alice B.")],
        year="2023", title="Nurse staffing and patient falls",
        container="Journal of Nursing Scholarship",
        volume="55", issue="4", pages="412-425", doi="10.1111/jnu.12345",
    )
    work.ensure_key()
    project = Project(
        topic="Nurse Staffing Ratios and Inpatient Fall Rates",
        works=[work],
        title_page=TitlePage(
            variant=variant,
            title="Nurse Staffing Ratios and Inpatient Fall Rates",
            authors=["Matthew Koch"],
            affiliations=["School of Nursing, Example University"],
            course="NURS 5100: Evidence-Based Practice",
            instructor="Dr. A. Reviewer",
            due_date="August 17, 2026",
            running_head=default_running_head(
                "Nurse Staffing Ratios and Inpatient Fall Rates"),
            author_note="Correspondence concerning this article should be addressed "
                        "to the first author.",
        ),
    )
    return project


def build(variant: str = "student", font: str = "Times New Roman") -> Path:
    project = sample_project(variant)
    context = CitationContext(project.works)
    paper = ApaPaper(project, context, font=font)
    paper.title_page()
    paper.abstract(
        "This review examined the association between nurse staffing ratios and "
        "inpatient fall rates across acute care settings.",
        keywords=["nurse staffing", "patient falls", "acute care"],
    )
    paper.body([
        Block("heading", level=1, text="Review of the Literature"),
        Block("paragraph", runs=[Run("Staffing levels are associated with falls "),
                                 Run("(Smith & Jones, 2023)"), Run(".")]),
        Block("heading", level=2, text="Measurement Approaches"),
        Block("paragraph", runs=[Run("Approaches differ across settings.")]),
        Block("heading", level=3, text="Unit-Level Measures"),
        Block("paragraph", runs=[Run("Unit-level measures predominate.")]),
        Block("heading", level=4, text="Hours per patient day"),
        Block("heading", level=5, text="Skill mix"),
        Block("quote", runs=[Run("A long quotation of forty or more words is set "
                                 "as an indented block without quotation marks.")]),
    ])
    paper.table(1, "Characteristics of Included Studies",
                ["Author", "Design", "n"],
                [["Smith & Jones (2023)", "Cohort", "1,204"]],
                note="n = analytic sample size.")
    paper.references(project.works)

    path = Path(tempfile.mkdtemp()) / f"paper-{variant}.docx"
    paper.save(str(path))
    return path


# ------------------------------------------------------------------ page setup


def test_margins_are_one_inch() -> None:
    doc = Document(str(build()))
    section = doc.sections[0]
    # APA 7 §2.22
    for name in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        check(f"{name} is 1 inch", getattr(section, name), Inches(1))


def test_font_set_on_all_four_ooxml_attributes() -> None:
    doc = Document(str(build()))
    rpr = doc.styles["Normal"].element.find(qn("w:rPr"))
    fonts = rpr.find(qn("w:rFonts"))
    # A font set only on w:ascii silently changes typeface at the first
    # non-Latin-1 character (APA 7 §2.19 applies to the whole document).
    for attribute in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        check(f"Normal style {attribute}", fonts.get(qn(attribute)), "Times New Roman")
    size = rpr.find(qn("w:sz"))
    check("Normal style size is 12 pt (24 half-points)", size.get(qn("w:val")), "24")


def test_rejects_font_apa_does_not_list() -> None:
    project = sample_project()
    context = CitationContext(project.works)
    try:
        ApaPaper(project, context, font="Comic Sans MS")
    except ValueError as error:
        check("unlisted font rejected", "APA 7 §2.19" in str(error), True)
    else:
        check("unlisted font rejected", False, True)


def test_calibri_defaults_to_eleven_point() -> None:
    doc = Document(str(build(font="Calibri")))
    rpr = doc.styles["Normal"].element.find(qn("w:rPr"))
    # §2.19 pairs each approved typeface with its own size.
    check("Calibri default size is 11 pt", rpr.find(qn("w:sz")).get(qn("w:val")), "22")


def test_double_spacing_throughout() -> None:
    doc = Document(str(build()))
    spacing = doc.styles["Normal"].element.find(qn("w:pPr")).find(qn("w:spacing"))
    # APA 7 §2.21. Word encodes double spacing as line=480 (240ths of a line).
    check("Normal line rule", spacing.get(qn("w:lineRule")), "auto")
    check("Normal line spacing", spacing.get(qn("w:line")), "480")

    singles = [
        p for p in doc.paragraphs
        if p.text.strip()
        and p.paragraph_format.line_spacing is not None
        and float(p.paragraph_format.line_spacing) < 2.0
    ]
    check("no single-spaced body paragraphs", singles, [])


# -------------------------------------------------------------------- header


def test_page_number_is_a_field_not_a_literal() -> None:
    doc = Document(str(build()))
    header_xml = doc.sections[0].header.paragraphs[0]._p.xml
    # A literal "1" would number every page 1. APA 7 §2.18 needs a real field.
    check("PAGE field instruction present", "PAGE" in header_xml, True)
    check("field begin char present", 'w:fldCharType="begin"' in header_xml, True)
    check("field end char present", 'w:fldCharType="end"' in header_xml, True)


def test_student_header_has_no_running_head() -> None:
    doc = Document(str(build("student")))
    header = doc.sections[0].header.paragraphs[0]
    # APA 7 §2.18: student papers carry the page number and nothing else.
    #
    # The header text is "1" rather than empty because the PAGE field carries a
    # cached result, so the number still shows in readers that do not evaluate
    # fields. What matters here is that no running head accompanies it.
    check("student header has no running head text",
          any(c.isalpha() for c in header.text), False)
    check("student header is right-aligned", header.alignment, 2)  # RIGHT
    check("student header still has the page field", "PAGE" in header._p.xml, True)


def test_professional_header_has_running_head() -> None:
    doc = Document(str(build("professional")))
    header = doc.sections[0].header.paragraphs[0]
    # APA 7 §2.8: capitals, flush left, 50 characters or fewer.
    check("running head present", "NURSE STAFFING" in header.text, True)
    check("running head is uppercase", header.text.strip().split("\t")[0].isupper(), True)
    check("running head within 50 chars", len(header.text.strip().split("\t")[0]) <= 50, True)


def test_header_appears_on_first_page_too() -> None:
    doc = Document(str(build()))
    # APA 7 §2.18: the title page is page 1 and carries the number.
    check("no separate first-page header",
          doc.sections[0].different_first_page_header_footer, False)


def test_running_head_truncates_on_word_boundary() -> None:
    long_title = ("Association Between Registered Nurse Staffing Ratios and "
                  "Inpatient Fall Rates in Acute Care Hospitals")
    head = default_running_head(long_title)
    check("truncated to 50 or fewer", len(head) <= 50, True)
    check("no partial trailing word", head.endswith(("S", "D", "N", "E", "O", "Y", "G")), True)
    check("uppercase", head.isupper(), True)


# ---------------------------------------------------------------- title page


def test_student_title_page_carries_course_block() -> None:
    doc = Document(str(build("student")))
    text = "\n".join(p.text for p in doc.paragraphs[:20])
    # APA 7 §2.3
    check("course present", "NURS 5100" in text, True)
    check("instructor present", "Dr. A. Reviewer" in text, True)
    check("due date present", "August 17, 2026" in text, True)
    check("no author note on student page", "Author Note" in text, False)


def test_professional_title_page_carries_author_note() -> None:
    doc = Document(str(build("professional")))
    text = "\n".join(p.text for p in doc.paragraphs[:30])
    # APA 7 §2.7
    check("author note heading present", "Author Note" in text, True)
    check("no course block on professional page", "NURS 5100" in text, False)


def test_title_is_bold_and_centred() -> None:
    doc = Document(str(build()))
    title = next(p for p in doc.paragraphs if "Nurse Staffing Ratios" in p.text)
    check("title bold", all(r.bold for r in title.runs if r.text.strip()), True)
    check("title centred", title.alignment, 1)  # WD_ALIGN_PARAGRAPH.CENTER


# ------------------------------------------------------------------ abstract


def test_abstract_starts_new_page_and_has_no_indent() -> None:
    paras = paragraphs_of(build())
    index = find_index(paras, lambda p: p.text.strip() == "Abstract")
    heading = paras[index]
    check("abstract on a new page", "pageBreakBefore" in heading._p.xml, True)
    check("abstract heading bold", all(r.bold for r in heading.runs if r.text.strip()), True)
    check("abstract heading centred", heading.alignment, 1)

    body = paras[index + 1]
    # APA 7 §2.9: the abstract is the one paragraph in the paper with no
    # first-line indent.
    check("abstract text not indented",
          body.paragraph_format.first_line_indent in (None, 0), True)


def test_keywords_label_is_italic_and_indented() -> None:
    doc = Document(str(build()))
    paragraph = next(p for p in doc.paragraphs if p.text.startswith("Keywords:"))
    # APA 7 §2.10
    check("Keywords label italic", paragraph.runs[0].italic, True)
    check("Keywords terms not italic", bool(paragraph.runs[1].italic), False)
    near("Keywords indented half an inch",
         paragraph.paragraph_format.first_line_indent, EMU_PER_INCH // 2)
    check("no terminal period after keywords", paragraph.text.rstrip().endswith("."), False)


# ------------------------------------------------------------------ headings


def test_five_heading_levels_are_distinct() -> None:
    doc = Document(str(build()))

    def find(text):
        return next(p for p in doc.paragraphs if p.text.strip().startswith(text))

    # APA 7 §2.27
    level1 = find("Review of the Literature")
    check("L1 centred", level1.alignment, 1)
    check("L1 bold", level1.runs[0].bold, True)
    check("L1 not italic", bool(level1.runs[0].italic), False)

    level2 = find("Measurement Approaches")
    check("L2 flush left", level2.alignment in (0, None), True)
    check("L2 bold", level2.runs[0].bold, True)
    check("L2 not italic", bool(level2.runs[0].italic), False)

    level3 = find("Unit-Level Measures")
    check("L3 bold", level3.runs[0].bold, True)
    check("L3 italic", level3.runs[0].italic, True)

    level4 = find("Hours per patient day")
    check("L4 bold", level4.runs[0].bold, True)
    check("L4 not italic", bool(level4.runs[0].italic), False)
    check("L4 ends with a period", level4.text.strip().endswith("."), True)
    near("L4 indented", level4.paragraph_format.first_line_indent, EMU_PER_INCH // 2)

    level5 = find("Skill mix")
    check("L5 bold", level5.runs[0].bold, True)
    check("L5 italic", level5.runs[0].italic, True)
    check("L5 ends with a period", level5.text.strip().endswith("."), True)


def test_headings_keep_with_next() -> None:
    doc = Document(str(build()))
    heading = next(p for p in doc.paragraphs if p.text.strip() == "Measurement Approaches")
    # A heading orphaned at the foot of a page is a layout defect Word will
    # produce unless told otherwise.
    check("heading keeps with next", "keepNext" in heading._p.xml, True)


# --------------------------------------------------------------- body & quotes


def test_body_paragraphs_are_indented() -> None:
    doc = Document(str(build()))
    body = next(p for p in doc.paragraphs
                if p.text.startswith("Staffing levels are associated"))
    # APA 7 §2.24
    near("body first-line indent", body.paragraph_format.first_line_indent,
         EMU_PER_INCH // 2)
    check("no space between paragraphs",
          body.paragraph_format.space_after in (None, 0), True)


def test_block_quote_is_indented_without_quote_marks() -> None:
    doc = Document(str(build()))
    quote = next(p for p in doc.paragraphs if p.text.startswith("A long quotation"))
    # APA 7 §8.27
    near("block quote left indent", quote.paragraph_format.left_indent,
         EMU_PER_INCH // 2)
    check("block quote has no quotation marks", '"' in quote.text or "“" in quote.text, False)


# ------------------------------------------------------------------ references


def test_references_start_new_page_with_hanging_indent() -> None:
    paras = paragraphs_of(build())
    index = find_index(paras, lambda p: p.text.strip() == "References")
    heading = paras[index]
    # APA 7 §2.12
    check("references on a new page", "pageBreakBefore" in heading._p.xml, True)
    check("references heading centred", heading.alignment, 1)
    check("references heading bold", heading.runs[0].bold, True)

    entry = paras[index + 1]
    # APA 7 §9.43: half-inch hanging indent, which in OOXML is a positive left
    # indent with a negative first-line indent of the same size.
    near("entry left indent", entry.paragraph_format.left_indent, EMU_PER_INCH // 2)
    near("entry hanging indent", entry.paragraph_format.first_line_indent,
         -(EMU_PER_INCH // 2))
    check("entry text present", entry.text.startswith("Smith, J. A., & Jones, A. B."), True)


def test_reference_journal_name_is_italic_in_the_docx() -> None:
    paras = paragraphs_of(build())
    entry = paras[find_index(paras, lambda p: p.text.strip() == "References") + 1]
    italic_text = "".join(r.text for r in entry.runs if r.italic)
    # The whole point of carrying styled runs through the pipeline: APA
    # italicises the journal name and volume and nothing else in the entry.
    check("journal name italic in file", "Journal of Nursing Scholarship" in italic_text, True)
    check("volume italic in file", italic_text.strip().endswith("55"), True)
    check("article title not italic", "Nurse staffing and patient falls" in italic_text, False)


# ---------------------------------------------------------------------- tables


def test_apa_table_has_horizontal_rules_only() -> None:
    doc = Document(str(build()))
    table = doc.tables[0]
    borders = table._tbl.tblPr.find(qn("w:tblBorders"))
    check("table borders element present", borders is not None, True)
    # APA 7 §7.8: no vertical lines, and no interior horizontal lines between
    # data rows. python-docx's "Table Grid" is the exact opposite.
    for edge in ("left", "right", "insideV", "insideH"):
        element = borders.find(qn(f"w:{edge}"))
        check(f"table {edge} border none", element.get(qn("w:val")), "none")
    for edge in ("top", "bottom"):
        element = borders.find(qn(f"w:{edge}"))
        check(f"table {edge} border single", element.get(qn("w:val")), "single")


def test_table_number_and_title_formatting() -> None:
    paras = paragraphs_of(build())
    index = find_index(paras, lambda p: p.text.strip() == "Table 1")
    check("table number bold", paras[index].runs[0].bold, True)
    title = paras[index + 1]
    # APA 7 §7.10: the title is italic Title Case on its own line.
    check("table title italic", title.runs[0].italic, True)
    check("table title text", title.text.strip(), "Characteristics of Included Studies")


def test_table_note_label_is_italic() -> None:
    doc = Document(str(build()))
    note = next(p for p in doc.paragraphs if p.text.startswith("Note."))
    # APA 7 §7.14
    check("Note. label italic", note.runs[0].italic, True)
    check("note body not italic", bool(note.runs[1].italic), False)


def test_header_row_repeats_across_pages() -> None:
    doc = Document(str(build()))
    tr_pr = doc.tables[0].rows[0]._tr.find(qn("w:trPr"))
    # APA 7 §7.13
    check("header row marked as repeating", tr_pr.find(qn("w:tblHeader")) is not None, True)


def main() -> int:
    tests = [value for name, value in sorted(globals().items())
             if name.startswith("test_") and callable(value)]
    for test in tests:
        test()
    print(f"APA 7 Word layout: {CHECKS} checks, {len(FAILURES)} failed")
    for failure in FAILURES:
        print(f"  FAIL {failure}")
    return 1 if FAILURES else 0


if __name__ == "__main__":
    raise SystemExit(main())
