"""
Export a shift's charting to a Word document.

Two documents come out of a shift, for the same reason the research half of this
suite produces a paper and a companion audit document: one is the work, the other
is the evidence about the work.

**The shift record** is the notes themselves, in chronological order, ready to be
read alongside the EHR while transcribing. Untranscribed notes are printed **first
and marked**, because the whole point of the document is to close that gap.

**The audit appendix** is the provenance: every timestamp pair, every revision with
its superseded text struck through, every override and its reason, every flag that
was acknowledged, the hash-chain verification, and the open reassessment loops. If
a note is ever questioned, this is the document that answers *when did you write
this, what did you change, and why*.

## Why revisions are struck through rather than replaced

The specification asked for non-destructive editing with strike-through and a
required addendum reason, and that is exactly what a paper chart does — you draw a
single line through the error, write the correction, initial and date it, and the
original stays legible. Reproducing that in the export is the difference between a
correction and an alteration. A record that shows what it used to say is
trustworthy; one that silently shows only the current text invites the question of
what it said before.

Every document carries `disclosure.EXPORT_HEADER` on its first page. Export is the
moment the text leaves the tool, so it is the moment the "not the legal medical
record" statement matters most.
"""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from ..apa import ooxml
from . import disclosure, interlocks, narrative as narrative_module, store as store_module
from .models import Encounter, Entry, Severity, parse_iso

# 11-point Calibri: this is a working clinical document rather than an APA
# manuscript, and it is read on a screen at speed. The APA font rules in
# `apa/document.py` do not apply here and pretending otherwise would make the
# document worse at its job.
BODY_FONT = "Calibri"
BODY_SIZE = 11.0


def _clock(value: str) -> str:
    stamp = parse_iso(value)
    return f"{stamp:%H:%M}" if stamp else "—"


def _stamp(value: str) -> str:
    moment = parse_iso(value)
    return f"{moment:%Y-%m-%d %H:%M} UTC" if moment else "—"


def _setup(document: Document) -> None:
    style = document.styles["Normal"]
    ooxml.force_style_font(style, BODY_FONT, BODY_SIZE)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    for section in document.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
    ooxml.set_document_language(document, "en-US")


def _heading(document: Document, text: str, size: float = 14.0,
             space_before: float = 12.0) -> Any:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    run.bold = True
    ooxml.force_font(run, BODY_FONT, size)
    ooxml.keep_with_next(paragraph)
    return paragraph


def _body(document: Document, text: str, *, italic: bool = False,
          size: float = BODY_SIZE, indent: float = 0.0) -> Any:
    paragraph = document.add_paragraph()
    if indent:
        paragraph.paragraph_format.left_indent = Inches(indent)
    for index, line in enumerate(str(text).split("\n")):
        run = paragraph.add_run(("" if index == 0 else "\n") + line)
        run.italic = italic
        ooxml.force_font(run, BODY_FONT, size)
    return paragraph


def _callout(document: Document, text: str, *, colour: tuple[int, int, int] | None = None) -> Any:
    """A bordered, indented statement — used for the disclosure blocks."""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.italic = True
    ooxml.force_font(run, BODY_FONT, 10.0)
    if colour:
        run.font.color.rgb = RGBColor(*colour)
    return paragraph


def _struck(document: Document, text: str) -> Any:
    """Superseded text, single-line struck through and still legible.

    python-docx exposes `font.strike`, so no raw OOXML is needed — unlike the page
    fields and table borders in `apa/ooxml.py`, this one the library covers.
    """
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    run = paragraph.add_run(text)
    run.font.strike = True
    ooxml.force_font(run, BODY_FONT, 10.0)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return paragraph


def _kv_table(document: Document, rows: list[tuple[str, str]],
              widths: tuple[float, float] = (2.0, 4.6)) -> Any:
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    ooxml.clear_cell_margins_shading(table)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(widths[0])
        cells[1].width = Inches(widths[1])
        left = cells[0].paragraphs[0].add_run(label)
        left.bold = True
        ooxml.force_font(left, BODY_FONT, 10.0)
        right = cells[1].paragraphs[0].add_run(value or "—")
        ooxml.force_font(right, BODY_FONT, 10.0)
    return table


# ------------------------------------------------------------------ shift record


def build_shift_record(encounter: Encounter, *,
                       author: str = "",
                       at: datetime | None = None) -> Document:
    document = Document()
    _setup(document)

    _heading(document, "Nursing documentation draft", 18.0, space_before=0.0)
    _callout(document, disclosure.EXPORT_HEADER, colour=(0xB0, 0x30, 0x00))

    _kv_table(document, [
        ("Encounter label", encounter.local_label or "unlabelled"),
        ("Initials", encounter.initials or "—"),
        ("Setting", encounter.setting.label),
        ("Age band", encounter.age_band or "—"),
        ("Weight", f"{encounter.weight_kg:g} kg" if encounter.weight_kg is not None
         else "not recorded"),
        ("Code status", encounter.code_status or "not recorded"),
        ("Allergies", encounter.allergies or "not recorded"),
        ("Isolation", encounter.isolation or "none recorded"),
        ("Language and interpreter",
         (f"{encounter.language}"
          + (f", interpreter {encounter.interpreter_id}"
             if encounter.interpreter_used else ", no interpreter recorded"))
         if encounter.language else "not recorded"),
        ("Notes in this document", str(len(encounter.entries))),
        ("Exported", _stamp((at or datetime.now(timezone.utc)).isoformat())),
        ("Exported by", author or "—"),
    ])

    untranscribed = encounter.untranscribed()
    if untranscribed:
        _heading(document, f"Not yet transcribed into the EHR — {len(untranscribed)} "
                           f"note{'s' if len(untranscribed) != 1 else ''}", 13.0)
        _callout(document,
                 "These are printed first because they are the open task. A note "
                 "that lives only here is a note that is not in the patient's "
                 "record, and a note here that disagrees with the record is worse "
                 "than no note at all. Transcribe them, then mark them "
                 "transcribed.",
                 colour=(0xB0, 0x30, 0x00))
        for entry in untranscribed:
            _render_entry(document, entry, encounter)

    transcribed = [e for e in encounter.sorted_entries() if e.transcribed_at]
    if transcribed:
        _heading(document, "Transcribed notes", 13.0)
        for entry in transcribed:
            _render_entry(document, entry, encounter)

    if not encounter.entries:
        _body(document, "No notes were composed on this encounter.", italic=True)

    _render_open_loops(document, encounter)
    _render_shift_context(document, encounter)
    return document


def _render_entry(document: Document, entry: Entry, encounter: Encounter) -> None:
    title = entry.title()
    _heading(document, title, 12.0)

    _kv_table(document, [
        ("Event time", _stamp(entry.event_at)),
        ("Documented", _stamp(entry.documented_at)
         + (f"  ({entry.late_by_minutes} min after the event)"
            if entry.late_by_minutes else "")),
        ("Author", " ".join(x for x in (entry.author_initials,
                                        entry.author_credential) if x) or "—"),
        ("Transcribed into the EHR",
         _stamp(entry.transcribed_at) if entry.transcribed_at
         else "NOT YET TRANSCRIBED"),
    ], widths=(1.8, 4.8))

    body = entry.narrative.strip() or narrative_module.compose(entry, encounter)
    _body(document, body)

    if entry.revisions:
        _heading(document, "Corrections to this note", 10.5, space_before=8.0)
        _body(document,
              "The superseded text is shown struck through and remains legible, "
              "as it would on a paper chart. Each correction carries the reason "
              "it was made.", italic=True, size=9.5)
        for revision in entry.revisions:
            _body(document,
                  f"{_stamp(revision.revised_at)} — "
                  f"{revision.author_initials or 'author not recorded'} — "
                  f"reason: {revision.reason}",
                  size=9.5, indent=0.25)
            _struck(document, revision.previous_text or "(no previous text)")

    acknowledged = [f for f in entry.flags
                    if f.get("code") in entry.acknowledged_warnings
                    or f.get("code") == "override_recorded"]
    if acknowledged:
        _heading(document, "Warnings shown and acknowledged", 10.5, space_before=8.0)
        _body(document,
              "Recorded because a warning that was seen and considered is "
              "stronger evidence of judgement than a warning that never fired.",
              italic=True, size=9.5)
        for flag in acknowledged:
            line = f"{flag.get('code')}: {flag.get('message', '')}"
            if flag.get("excerpt"):
                line += f"  [{flag['excerpt']}]"
            _body(document, line, size=9.5, indent=0.25)


def _render_open_loops(document: Document, encounter: Encounter) -> None:
    loops = encounter.reassessments
    if not loops:
        return
    _heading(document, "Reassessments", 13.0)
    table = document.add_table(rows=1, cols=4)
    ooxml.apa_table_borders(table)
    headers = ["Owed for", "Due", "Performed", "Finding"]
    for index, text in enumerate(headers):
        run = table.rows[0].cells[index].paragraphs[0].add_run(text)
        run.bold = True
        ooxml.force_font(run, BODY_FONT, 10.0)
    for loop in loops:
        cells = table.add_row().cells
        status = (loop.finding if loop.satisfied
                  else (f"not done: {loop.not_done_reason}" if loop.not_done_reason
                        else ("OVERDUE" if loop.is_overdue() else "open")))
        values = [loop.trigger, _clock(loop.due_at),
                  _clock(loop.performed_at) if loop.performed_at else "—", status]
        for index, value in enumerate(values):
            run = cells[index].paragraphs[0].add_run(value)
            ooxml.force_font(run, BODY_FONT, 9.5)


def _render_shift_context(document: Document, encounter: Encounter) -> None:
    context = encounter.context
    if not any([context.unit, context.patient_count, context.concerns_summary,
                context.expected_ratio, context.acuity_note]):
        return
    _heading(document, "Assignment and staffing record", 13.0)
    _callout(document,
             "Kept separately from the clinical notes on purpose. A "
             "contemporaneous record of the conditions of care is strong evidence "
             "about the circumstances; the same information inside a patient's "
             "chart reads as an excuse and gives the other side a theory. This "
             "page is about the shift, not the patient.")
    _kv_table(document, [
        ("Unit", context.unit),
        ("Shift", context.shift),
        ("Patients assigned",
         str(context.patient_count) if context.patient_count is not None else ""),
        ("Expected ratio", context.expected_ratio),
        ("Acuity factors", context.acuity_note),
        ("Support available", context.support_staff),
        ("Charge nurse", context.charge_nurse),
        ("Concern raised with", context.concerns_raised_to),
        ("Time raised", _clock(context.concerns_raised_at)),
        ("Concern as stated", context.concerns_summary),
        ("Response received", context.response_received),
    ])


# ---------------------------------------------------------------- audit appendix


def build_audit_appendix(encounter: Encounter, *,
                        author: str = "",
                        at: datetime | None = None) -> Document:
    document = Document()
    _setup(document)

    _heading(document, "Documentation audit appendix", 18.0, space_before=0.0)
    _callout(document, disclosure.EXPORT_HEADER, colour=(0xB0, 0x30, 0x00))
    _body(document,
          "This document is the provenance of the shift record: when each note "
          "was written against when its event occurred, every correction with its "
          "superseded text and reason, every interlock that fired, and the "
          "verification of the audit chain. It exists so that the question \"when "
          "did you write this, and what did you change\" has a written answer "
          "rather than a recollection.")

    # 1. Timestamp ledger. The whole point is the pair, side by side.
    _heading(document, "1. Timestamp ledger", 14.0)
    _body(document,
          "Documentation time comes from the system clock and cannot be set by "
          "any user of this tool. Event time is entered by the nurse. Where the "
          "gap exceeds the late-entry window the note is tagged [LATE ENTRY] "
          "automatically — which is the correct way to document something "
          "recalled later, and reads far better than a timestamp that does not "
          "match the medication scanner.", italic=True, size=10.0)
    table = document.add_table(rows=1, cols=5)
    ooxml.apa_table_borders(table)
    for index, text in enumerate(["Note", "Event time", "Documented",
                                  "Gap (min)", "Late entry"]):
        run = table.rows[0].cells[index].paragraphs[0].add_run(text)
        run.bold = True
        ooxml.force_font(run, BODY_FONT, 10.0)
    for entry in encounter.sorted_entries():
        cells = table.add_row().cells
        values = [entry.kind.label, _stamp(entry.event_at),
                  _stamp(entry.documented_at), str(entry.late_by_minutes),
                  "yes" if entry.late_entry else "no"]
        for index, value in enumerate(values):
            run = cells[index].paragraphs[0].add_run(value)
            ooxml.force_font(run, BODY_FONT, 9.5)

    # 2. Corrections.
    _heading(document, "2. Corrections", 14.0)
    revised = [e for e in encounter.sorted_entries() if e.revisions]
    if not revised:
        _body(document, "No note on this encounter has been corrected.",
              italic=True)
    else:
        _body(document,
              "Nothing in this tool overwrites a note. Each correction below "
              "keeps the text it replaced, struck through and legible, exactly as "
              "a single line through an error on a paper chart would.",
              italic=True, size=10.0)
        for entry in revised:
            _heading(document, entry.title(), 11.5, space_before=10.0)
            for revision in entry.revisions:
                _body(document,
                      f"{_stamp(revision.revised_at)} — "
                      f"{revision.author_initials or 'author not recorded'} — "
                      f"reason: {revision.reason}", size=10.0, indent=0.25)
                _body(document, "Superseded text:", size=9.5, indent=0.25,
                      italic=True)
                _struck(document, revision.previous_text or "(no previous text)")
                _body(document, "Replaced with:", size=9.5, indent=0.25,
                      italic=True)
                _body(document, revision.new_text, size=10.0, indent=0.4)

    # 3. Interlocks and overrides.
    _heading(document, "3. Interlocks, warnings and overrides", 14.0)
    _body(document,
          "What the tool objected to, and what was done about it. An override "
          "with a stated reason is a stronger artefact than a block would have "
          "been: it shows a person making a decision under pressure and saying "
          "why, which is what actually happened.", italic=True, size=10.0)
    any_flags = False
    for entry in encounter.sorted_entries():
        if not entry.flags:
            continue
        any_flags = True
        _heading(document, entry.title(), 11.5, space_before=10.0)
        for flag in entry.flags:
            severity = str(flag.get("severity", "")).upper()
            line = f"[{severity}] {flag.get('code', '')}: {flag.get('message', '')}"
            if flag.get("excerpt"):
                line += f"  Text flagged: “{flag['excerpt']}”"
            if flag.get("code") in entry.acknowledged_warnings:
                line += "  — acknowledged by the author."
            _body(document, line, size=9.5, indent=0.25)
    if not any_flags:
        _body(document, "No interlock fired on any note on this encounter.",
              italic=True)

    # 4. Scoring instruments.
    _heading(document, "4. Scoring instruments used", 14.0)
    used: dict[str, Any] = {}
    for entry in encounter.sorted_entries():
        for scale in entry.scales:
            used.setdefault(scale.scale_id, scale)
    if not used:
        _body(document, "No scoring instrument was recorded on this encounter.",
              italic=True)
    else:
        _body(document,
              "Each score below was computed from selections the nurse made. The "
              "instruments are cited because most of them are copyrighted works "
              "whose item wording is not reproduced in this tool — the scoring "
              "logic and the published thresholds are implemented, and the "
              "authoritative definitions live in your unit's licensed copy.",
              italic=True, size=10.0)
        for scale in used.values():
            _body(document, f"{scale.scale_name} — {scale.citation}",
                  size=9.5, indent=0.25)

    # 5. Chain verification.
    _heading(document, "5. Audit chain verification", 14.0)
    result = store_module.verify_chain(encounter)
    _kv_table(document, [
        ("Events recorded", str(len(encounter.audit))),
        ("Chain intact", "yes" if result["intact"] else "NO"),
        ("Events verified", str(result["verified"])),
        ("First break",
         "none" if result["intact"] else f"event {result['broken_at']} "
                                        f"({result['event_id']})"),
    ])
    _body(document, result["reason"], size=10.0)

    table = document.add_table(rows=1, cols=5)
    ooxml.apa_table_borders(table)
    for index, text in enumerate(["#", "Time", "Action", "Target", "Digest"]):
        run = table.rows[0].cells[index].paragraphs[0].add_run(text)
        run.bold = True
        ooxml.force_font(run, BODY_FONT, 9.5)
    for index, event in enumerate(encounter.audit):
        cells = table.add_row().cells
        values = [str(index), _stamp(event.at), event.action,
                  event.target or "—", event.digest[:12] + "…"]
        for position, value in enumerate(values):
            run = cells[position].paragraphs[0].add_run(value)
            ooxml.force_font(run, BODY_FONT, 8.5)

    # 6. What this document cannot do.
    _heading(document, "6. What this record does not establish", 14.0)
    for item in disclosure.what_it_cannot_do():
        paragraph = document.add_paragraph()
        run = paragraph.add_run(item["claim"] + ". ")
        run.bold = True
        ooxml.force_font(run, BODY_FONT, 10.0)
        tail = paragraph.add_run(item["reality"])
        ooxml.force_font(tail, BODY_FONT, 10.0)

    return document


# ---------------------------------------------------------------------- writing


def write_documents(encounter: Encounter, directory: Path, *,
                    author: str = "") -> list[dict[str, str]]:
    """Write both documents and return what was produced."""
    directory.mkdir(parents=True, exist_ok=True)
    label = (encounter.local_label or encounter.encounter_id).replace(" ", "-")
    safe = "".join(c for c in label if c.isalnum() or c in "-_")[:40] or "encounter"

    produced: list[dict[str, str]] = []
    record_path = directory / f"{safe}-shift-record.docx"
    build_shift_record(encounter, author=author).save(str(record_path))
    produced.append({
        "name": record_path.name,
        "kind": "shift record",
        "description": ("The notes themselves, chronologically, with untranscribed "
                        "notes first."),
    })

    audit_path = directory / f"{safe}-audit-appendix.docx"
    build_audit_appendix(encounter, author=author).save(str(audit_path))
    produced.append({
        "name": audit_path.name,
        "kind": "audit appendix",
        "description": ("Timestamp ledger, corrections with superseded text, "
                        "interlocks and overrides, instruments cited, and the "
                        "audit chain verification."),
    })
    return produced
