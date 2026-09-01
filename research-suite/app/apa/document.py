"""
The APA 7th edition Word paper.

Layout rules implemented here, each cited to the *Publication Manual* (7th ed.):

* 1-inch margins on all sides (§2.22)
* double spacing throughout — body, block quotations, table notes, and the
  reference list alike (§2.21)
* 0.5-inch first-line indent on body paragraphs, and no extra space between
  them (§2.24)
* page number flush right in the header of every page, title page included
  (§2.18)
* a running head in capitals for professional papers only (§2.8)
* five heading levels with distinct formatting (§2.27)
* reference list on a new page with a 0.5-inch hanging indent (§2.12, §9.43)
* block quotations of 40 or more words indented without quotation marks (§8.27)

The builder takes `list[Run]` paragraphs rather than strings, because APA
italicises specific spans — a journal name in a reference, a genus name in body
text, the word "Keywords" — and a string-based API cannot express that.
"""

from __future__ import annotations

from dataclasses import dataclass, field

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt

from ..models import Project, TitlePage, Work
from . import ooxml
from .citations import CitationContext, Run, reference_list

# APA 7 §2.19 accepts several typefaces. Times New Roman 12 pt is the default
# here because it is what most nursing and public-health programmes specify.
APPROVED_FONTS = {
    "Times New Roman": 12.0,
    "Calibri": 11.0,
    "Arial": 11.0,
    "Georgia": 11.0,
    "Lucida Sans Unicode": 10.0,
    "Computer Modern": 10.0,
}

BODY_INDENT = Inches(0.5)
HANGING_INDENT = Inches(0.5)
BLOCK_QUOTE_INDENT = Inches(0.5)
BLOCK_QUOTE_MIN_WORDS = 40


@dataclass
class Block:
    """One rendered element of the paper body."""

    kind: str                       # paragraph | heading | quote | list-item
    level: int = 0                  # heading level 1-5
    runs: list[Run] = field(default_factory=list)
    text: str = ""                  # convenience for headings


class ApaPaper:
    """Builds an APA 7 paper into a python-docx Document."""

    def __init__(
        self,
        project: Project,
        context: CitationContext,
        font: str = "Times New Roman",
        font_size: float | None = None,
    ):
        if font not in APPROVED_FONTS:
            raise ValueError(
                f"{font!r} is not one of the typefaces APA 7 §2.19 lists. "
                f"Choose from: {', '.join(sorted(APPROVED_FONTS))}."
            )
        self.project = project
        self.context = context
        self.font = font
        self.font_size = font_size if font_size is not None else APPROVED_FONTS[font]
        self.document = Document()
        self._configure()

    # ------------------------------------------------------------ page setup

    def _configure(self) -> None:
        section = self.document.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        normal = self.document.styles["Normal"]
        ooxml.force_style_font(normal, self.font, self.font_size)
        fmt = normal.paragraph_format
        fmt.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.widow_control = True

        page = self.project.title_page
        running = page.running_head if page.variant == "professional" else ""
        ooxml.build_apa_header(section, running)
        ooxml.set_document_language(self.document)

    # ------------------------------------------------------------- primitives

    def _paragraph(
        self,
        runs: list[Run] | None = None,
        *,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=None,
        left_indent=None,
        bold: bool = False,
        italic: bool = False,
        keep_next: bool = False,
        new_page: bool = False,
    ):
        paragraph = self.document.add_paragraph()
        fmt = paragraph.paragraph_format
        fmt.alignment = align
        fmt.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        if first_line_indent is not None:
            fmt.first_line_indent = first_line_indent
        if left_indent is not None:
            fmt.left_indent = left_indent
        if new_page:
            ooxml.page_break_before(paragraph)
        if keep_next:
            ooxml.keep_with_next(paragraph)
            ooxml.keep_lines_together(paragraph)
        ooxml.widow_control(paragraph)
        ooxml.suppress_auto_hyphenation(paragraph)

        for run_spec in runs or []:
            run = paragraph.add_run(run_spec.text)
            run.italic = run_spec.italic or italic
            run.bold = bold
            ooxml.force_font(run, self.font, self.font_size)
        return paragraph

    def _blank_line(self) -> None:
        self._paragraph([Run("")])

    # ------------------------------------------------------------ title page

    def title_page(self) -> None:
        """APA 7 §2.3 (student) and §2.2 (professional).

        Both put the title in bold Title Case in the upper half of the page.
        The student version then carries the course, instructor and due date;
        the professional version carries an author note in the lower half.
        """
        page = self.project.title_page
        # Three or four blank double-spaced lines put the title in the upper
        # half of the page, which is what §2.3 asks for rather than a fixed
        # measurement.
        for _ in range(3):
            self._blank_line()

        self._paragraph(
            [Run(page.title or self.project.topic)],
            align=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
        )
        self._blank_line()

        for line in page.authors or [""]:
            self._paragraph([Run(line)], align=WD_ALIGN_PARAGRAPH.CENTER)
        for line in page.affiliations:
            self._paragraph([Run(line)], align=WD_ALIGN_PARAGRAPH.CENTER)

        if page.variant == "student":
            for line in (page.course, page.instructor, page.due_date):
                if line:
                    self._paragraph([Run(line)], align=WD_ALIGN_PARAGRAPH.CENTER)
        elif page.author_note:
            # §2.7: the author note sits in the bottom half, under a centred
            # bold heading, with flush-left indented paragraphs.
            for _ in range(6):
                self._blank_line()
            self._paragraph([Run("Author Note")], align=WD_ALIGN_PARAGRAPH.CENTER,
                            bold=True, keep_next=True)
            for para in page.author_note.split("\n\n"):
                if para.strip():
                    self._paragraph([Run(para.strip())], first_line_indent=BODY_INDENT)

    # -------------------------------------------------------------- abstract

    def abstract(self, text: str, keywords: list[str] | None = None) -> None:
        """APA 7 §2.9: own page, centred bold heading, and — unlike every other
        paragraph in the paper — no first-line indent."""
        if not text.strip():
            return
        heading = self._paragraph(
            [Run("Abstract")], align=WD_ALIGN_PARAGRAPH.CENTER, bold=True,
            keep_next=True, new_page=True,
        )
        del heading
        self._paragraph([Run(" ".join(text.split()))])

        if keywords:
            # §2.10: indented, "Keywords:" italic, terms lowercase and
            # comma-separated with no terminal period.
            paragraph = self._paragraph(first_line_indent=BODY_INDENT)
            label = paragraph.add_run("Keywords:")
            label.italic = True
            ooxml.force_font(label, self.font, self.font_size)
            body = paragraph.add_run(" " + ", ".join(k.strip() for k in keywords if k.strip()))
            ooxml.force_font(body, self.font, self.font_size)

    # ------------------------------------------------------------------ body

    def body(self, blocks: list[Block], *, title: str = "") -> None:
        """The paper body, starting on a new page with the title repeated
        centred and bold (APA 7 §2.11)."""
        page_title = title or self.project.title_page.title or self.project.topic
        self._paragraph(
            [Run(page_title)], align=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True, keep_next=True, new_page=True,
        )

        for block in blocks:
            if block.kind == "heading":
                self.heading(block.text or "", block.level)
            elif block.kind == "quote":
                self.block_quote(block.runs)
            elif block.kind == "list-item":
                self._paragraph(block.runs, left_indent=BODY_INDENT)
            else:
                self._paragraph(block.runs, first_line_indent=BODY_INDENT)

    def heading(self, text: str, level: int) -> None:
        """APA 7 §2.27, the five levels.

        Levels 4 and 5 are *run-in*: the heading ends with a period and the
        paragraph text continues on the same line. This builder emits them as
        their own paragraph with the period, because run-in headings are only
        correct when there is body text to run into — and the caller supplies
        text and heading separately. Callers wanting a true run-in should use
        `run_in_heading`.
        """
        level = max(1, min(5, level))
        if level == 1:
            self._paragraph([Run(text)], align=WD_ALIGN_PARAGRAPH.CENTER,
                            bold=True, keep_next=True)
        elif level == 2:
            self._paragraph([Run(text)], bold=True, keep_next=True)
        elif level == 3:
            self._paragraph([Run(text)], bold=True, italic=True, keep_next=True)
        elif level == 4:
            self._paragraph([Run(_ends_with_period(text))], bold=True,
                            first_line_indent=BODY_INDENT, keep_next=True)
        else:
            self._paragraph([Run(_ends_with_period(text))], bold=True, italic=True,
                            first_line_indent=BODY_INDENT, keep_next=True)

    def run_in_heading(self, text: str, level: int, runs: list[Run]) -> None:
        """A Level 4 or 5 heading with its paragraph on the same line."""
        paragraph = self._paragraph(first_line_indent=BODY_INDENT)
        head = paragraph.add_run(_ends_with_period(text) + " ")
        head.bold = True
        head.italic = level >= 5
        ooxml.force_font(head, self.font, self.font_size)
        for spec in runs:
            run = paragraph.add_run(spec.text)
            run.italic = spec.italic
            ooxml.force_font(run, self.font, self.font_size)

    def block_quote(self, runs: list[Run]) -> None:
        """APA 7 §8.27: quotations of 40 or more words are set as a block,
        indented half an inch, with no quotation marks."""
        self._paragraph(runs, left_indent=BLOCK_QUOTE_INDENT)

    # ------------------------------------------------------------ references

    def references(self, works: list[Work]) -> None:
        """APA 7 §2.12 and §9.43: new page, centred bold heading, every entry
        with a half-inch hanging indent, double spaced."""
        self._paragraph(
            [Run("References")], align=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True, keep_next=True, new_page=True,
        )
        if not works:
            self._paragraph([Run("No sources are cited in this draft.")],
                            first_line_indent=BODY_INDENT)
            return
        for entry in reference_list(works, self.context):
            paragraph = self._paragraph(entry)
            fmt = paragraph.paragraph_format
            fmt.left_indent = HANGING_INDENT
            fmt.first_line_indent = -HANGING_INDENT

    # ---------------------------------------------------------------- tables

    def table(
        self,
        number: int,
        title: str,
        headers: list[str],
        rows: list[list[str]],
        note: str = "",
    ) -> None:
        """An APA 7 table (§7.8-7.21).

        The number is plain and on its own line; the title is *italic* Title
        Case on the next line; the body carries horizontal rules only; a
        general note goes below, introduced by an italic "Note."
        """
        self._paragraph([Run(f"Table {number}")], bold=True, keep_next=True)
        self._paragraph([Run(title, italic=True)], keep_next=True)

        table = self.document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"          # replaced immediately below
        ooxml.apa_table_borders(table)
        ooxml.clear_cell_margins_shading(table)

        for cell, text in zip(table.rows[0].cells, headers):
            self._fill_cell(cell, text, bold=True)
        ooxml.rule_below_row(table.rows[0])

        for values in rows:
            row = table.add_row()
            for cell, text in zip(row.cells, values):
                self._fill_cell(cell, str(text))

        if note:
            paragraph = self._paragraph()
            label = paragraph.add_run("Note. ")
            label.italic = True
            ooxml.force_font(label, self.font, self.font_size)
            body = paragraph.add_run(note)
            ooxml.force_font(body, self.font, self.font_size)
        self._blank_line()

    def _fill_cell(self, cell, text: str, bold: bool = False) -> None:
        paragraph = cell.paragraphs[0]
        paragraph.text = ""
        fmt = paragraph.paragraph_format
        fmt.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        fmt.space_after = Pt(0)
        run = paragraph.add_run(text)
        run.bold = bold
        ooxml.force_font(run, self.font, self.font_size)

    # --------------------------------------------------------------- figures

    def figure(self, number: int, title: str, image_path: str, note: str = "",
               width_inches: float = 6.0) -> None:
        """An APA 7 figure (§7.22-7.36): bold number, italic title, image, then
        an italic-labelled note."""
        self._paragraph([Run(f"Figure {number}")], bold=True, keep_next=True)
        self._paragraph([Run(title, italic=True)], keep_next=True)
        self.document.add_picture(image_path, width=Inches(width_inches))
        # add_picture appends its own paragraph; centre it.
        self.document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
        if note:
            paragraph = self._paragraph()
            label = paragraph.add_run("Note. ")
            label.italic = True
            ooxml.force_font(label, self.font, self.font_size)
            body = paragraph.add_run(note)
            ooxml.force_font(body, self.font, self.font_size)
        self._blank_line()

    # --------------------------------------------------------------- appendix

    def appendix(self, label: str, title: str) -> None:
        """APA 7 §2.14: each appendix starts a new page with a centred bold
        label, then a centred bold title."""
        self._paragraph([Run(label)], align=WD_ALIGN_PARAGRAPH.CENTER,
                        bold=True, keep_next=True, new_page=True)
        if title:
            self._paragraph([Run(title)], align=WD_ALIGN_PARAGRAPH.CENTER,
                            bold=True, keep_next=True)

    # ------------------------------------------------------------------ save

    def save(self, path: str) -> str:
        self.document.save(path)
        return path


def _ends_with_period(text: str) -> str:
    text = text.strip()
    return text if text.endswith((".", "?", "!")) else text + "."


def is_block_quote(text: str) -> bool:
    """APA 7 §8.27: 40 or more words goes in a block."""
    return len(text.split()) >= BLOCK_QUOTE_MIN_WORDS


def default_running_head(title: str) -> str:
    """A running head APA will accept: capitals, 50 characters or fewer,
    truncated at a word boundary rather than mid-word (§2.8)."""
    head = " ".join(title.upper().split())
    if len(head) <= 50:
        return head
    cut = head[:50]
    if " " in cut:
        cut = cut[: cut.rfind(" ")]
    return cut


def new_title_page(project: Project, variant: str = "student") -> TitlePage:
    page = TitlePage(variant=variant, title=project.topic)
    page.running_head = default_running_head(page.title)
    return page


def make_section_break(document: Document) -> None:
    """A continuous section break — used by the audit document to switch to
    landscape for the evidence matrix without disturbing earlier pages."""
    document.add_section(WD_SECTION.NEW_PAGE)
