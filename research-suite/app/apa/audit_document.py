"""
The companion Rationale and Source Mapping document.

This is the second deliverable, and the one that makes the first defensible. It
answers, for a faculty reader, an IRB or a peer reviewer, every question they
would otherwise have to take on trust:

1. **What was searched, and how.** Every database, the exact query string, when
   it ran, how many records it returned. A search that cannot be re-run is not a
   method.
2. **What was screened out, and why.** With the inclusion and exclusion criteria
   stated first, and the PRISMA counts derived from the actual logs rather than
   asserted.
3. **The evidence matrix.** Author, year, design, sample, findings, strengths and
   limitations, level of evidence — one row per included study, on a landscape
   page so it is legible.
4. **The appraisal of each study**, with the instrument named and the rating's
   derivation shown.
5. **The claim-to-source map.** Every sentence of the paper, the source behind
   it, the locus within that source, the passage relied on, and why that source
   was chosen. This is the traceability matrix.
6. **What was excluded and why**, including near-misses — the studies a reviewer
   would ask about.
7. **Integrity checks.** Retraction checks with their negative results recorded,
   and the verbatim-overlap report.
8. **An AI-use disclosure** describing exactly which steps were assisted.

The document is formatted to APA conventions for readability but is not itself an
APA manuscript — it is an appendix-style working record, and it says so on its
title page so nobody mistakes it for the submission.
"""

from __future__ import annotations

from datetime import date

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt

from ..evidence import appraisal as appraisal_module
from ..evidence import dedupe as dedupe_module
from ..models import Claim, Project, SupportType, Work
from ..writing import integrity as integrity_module
from . import ooxml
from .citations import CitationContext, Run, plain, reference, reference_list


class AuditDocument:
    """Builds the companion document."""

    def __init__(
        self,
        project: Project,
        context: CitationContext,
        font: str = "Times New Roman",
        font_size: float = 12.0,
    ):
        self.project = project
        self.context = context
        self.font = font
        self.font_size = font_size
        self.document = Document()
        self._configure()

    def _configure(self) -> None:
        section = self.document.sections[0]
        for name in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
            setattr(section, name, Inches(1))
        normal = self.document.styles["Normal"]
        ooxml.force_style_font(normal, self.font, self.font_size)
        fmt = normal.paragraph_format
        # Single-spaced: this is a reference document to be read in tables and
        # scanned, not a manuscript to be marked up.
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(6)
        ooxml.build_apa_header(section, "")
        ooxml.set_document_language(self.document)

    # ------------------------------------------------------------ primitives

    def _para(self, text: str = "", *, bold=False, italic=False, size=None,
              align=WD_ALIGN_PARAGRAPH.LEFT, indent=None, space_after=6,
              keep_next=False, new_page=False, runs: list[Run] | None = None):
        paragraph = self.document.add_paragraph()
        fmt = paragraph.paragraph_format
        fmt.alignment = align
        fmt.space_after = Pt(space_after)
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
        if indent is not None:
            fmt.left_indent = indent
        if new_page:
            ooxml.page_break_before(paragraph)
        if keep_next:
            ooxml.keep_with_next(paragraph)
        for spec in (runs if runs is not None else [Run(text)]):
            if not spec.text:
                continue
            run = paragraph.add_run(spec.text)
            run.bold = bold
            run.italic = spec.italic or italic
            ooxml.force_font(run, self.font, size or self.font_size)
        return paragraph

    def _h1(self, text: str, *, new_page: bool = True) -> None:
        self._para(text, bold=True, size=self.font_size + 2, space_after=10,
                   keep_next=True, new_page=new_page)

    def _h2(self, text: str) -> None:
        self._para(text, bold=True, space_after=6, keep_next=True)

    def _bullet(self, text: str) -> None:
        self._para(f"•  {text}", indent=Inches(0.25), space_after=3)

    def _table(self, headers: list[str], rows: list[list[str]],
               widths: list[float] | None = None, font_size: float | None = None):
        table = self.document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        ooxml.apa_table_borders(table)
        ooxml.clear_cell_margins_shading(table)
        size = font_size or (self.font_size - 2)

        for cell, text in zip(table.rows[0].cells, headers):
            self._cell(cell, text, bold=True, size=size)
        ooxml.rule_below_row(table.rows[0])

        for values in rows:
            row = table.add_row()
            for cell, text in zip(row.cells, values):
                self._cell(cell, str(text), size=size)

        if widths:
            # Cell widths must be set per cell, not per column — Word ignores
            # column-level widths in many documents.
            for row in table.rows:
                for cell, width in zip(row.cells, widths):
                    cell.width = Inches(width)
        return table

    def _cell(self, cell, text: str, *, bold=False, size=None) -> None:
        paragraph = cell.paragraphs[0]
        paragraph.text = ""
        fmt = paragraph.paragraph_format
        fmt.space_after = Pt(2)
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = paragraph.add_run(text)
        run.bold = bold
        ooxml.force_font(run, self.font, size or (self.font_size - 2))

    def _landscape(self) -> None:
        section = self.document.add_section(WD_SECTION.NEW_PAGE)
        # Orientation alone does not rotate the page — the width and height must
        # be swapped as well, which is the step usually missed.
        width, height = section.page_width, section.page_height
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = height, width
        for name in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
            setattr(section, name, Inches(0.7))
        ooxml.build_apa_header(section, "")

    def _portrait(self) -> None:
        section = self.document.add_section(WD_SECTION.NEW_PAGE)
        width, height = section.page_width, section.page_height
        section.orientation = WD_ORIENT.PORTRAIT
        if width > height:
            section.page_width, section.page_height = height, width
        for name in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
            setattr(section, name, Inches(1))
        ooxml.build_apa_header(section, "")

    # ---------------------------------------------------------------- content

    def build(self, *, drafted: bool = True) -> "AuditDocument":
        self._title_page(drafted)
        self._search_record()
        self._screening_record()
        self._evidence_matrix()
        self._appraisal_summary()
        self._claim_map()
        self._exclusion_rationale()
        self._integrity_section()
        self._disclosure(drafted)
        self._references()
        return self

    # ------------------------------------------------------------ title page

    def _title_page(self, drafted: bool) -> None:
        page = self.project.title_page
        self._para(space_after=0)
        self._para(space_after=0)
        self._para("Rationale and Source Mapping Document", bold=True,
                   size=self.font_size + 4, align=WD_ALIGN_PARAGRAPH.CENTER,
                   space_after=12)
        self._para("Companion record for:", align=WD_ALIGN_PARAGRAPH.CENTER,
                   space_after=4)
        self._para(page.title or self.project.topic, italic=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

        for line in page.authors:
            self._para(line, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        for line in page.affiliations:
            self._para(line, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        if page.variant == "student":
            for line in (page.course, page.instructor, page.due_date):
                if line:
                    self._para(line, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        self._para(f"Prepared {date.today().isoformat()}",
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

        # Said plainly, on the first page, so this document is never mistaken for
        # the submission itself.
        self._para(
            "This document is a working record, not a manuscript. It is formatted "
            "for legibility rather than to APA manuscript specification, and is "
            "intended to accompany the paper so that a reader can trace every "
            "claim to its source. The paper itself is the APA 7 document.",
            italic=True, space_after=10,
        )
        self._para("Contents", bold=True, space_after=4)
        for index, name in enumerate([
            "How the literature was searched",
            "Screening decisions and record counts",
            "Evidence matrix of included studies",
            "Critical appraisal of each included study",
            "Claim-to-source traceability map",
            "Rationale for inclusion and exclusion",
            "Retraction and originality checks",
            "Disclosure of artificial intelligence use",
            "References",
        ], 1):
            self._para(f"{index}.  {name}", indent=Inches(0.25), space_after=2)

    # --------------------------------------------------------------- searches

    def _search_record(self) -> None:
        self._h1("1.  How the literature was searched")

        question = self.project.question or self.project.topic
        self._para(f"Research question: {question}", space_after=8)
        if self.project.pico:
            self._h2("PICO framing")
            for key, value in self.project.pico.items():
                if value:
                    self._bullet(f"{key.title()}: {value}")

        self._h2("Databases searched programmatically")
        if not self.project.searches:
            self._para("No automated searches are recorded for this project.")
        else:
            rows = [[
                s.get("source", ""),
                s.get("query", "")[:400],
                s.get("run_at", "")[:10],
                s.get("total_available", "—"),
                s.get("retrieved", "—"),
                "cache" if s.get("from_cache") == "yes" else "live",
            ] for s in self.project.searches]
            self._table(
                ["Database", "Exact query", "Date run", "Hits", "Retrieved", "Source"],
                rows, widths=[1.1, 2.6, 0.7, 0.6, 0.7, 0.6],
            )
            errors = [s for s in self.project.searches if s.get("error")]
            if errors:
                self._para("")
                self._para(
                    "Searches that failed and returned nothing (recorded so the "
                    "coverage of this review is not overstated):", bold=True)
                for search in errors:
                    self._bullet(f"{search.get('source')}: {search.get('error')}")
            notes = [s for s in self.project.searches if s.get("coverage_note")]
            for search in notes:
                self._para("")
                self._para(f"Coverage limit — {search.get('source')}: "
                           f"{search.get('coverage_note')}", italic=True)

        self._h2("Databases searched by hand and imported")
        imported = sorted({
            w.source_db for w in self.project.works if "(export)" in w.source_db
        })
        if imported:
            self._para(
                "Records from these databases were retrieved by running the search "
                "in the database's own interface and importing the citation "
                "export. They carry the same screening, appraisal and formatting "
                "as programmatically retrieved records; the difference in "
                "retrieval route is recorded here because it affects "
                "reproducibility — the query string is not captured "
                "automatically.", space_after=6)
            for name in imported:
                count = sum(1 for w in self.project.works if w.source_db == name)
                self._bullet(f"{name} — {count} record(s) imported")
        else:
            self._para("None.")

        self._h2("Databases with no route available to this project")
        self._para(
            "These are named explicitly so that the limits of this review's "
            "coverage are on the record. None of them sells API access to an "
            "unaffiliated individual, so records from them can only enter this "
            "project through a manual search and a citation export.",
            space_after=6)
        for name in ("CINAHL (EBSCOhost)", "PsycINFO", "Embase", "Scopus",
                     "Cochrane Library full text and CENTRAL",
                     "JBI EBP Database", "Global Health (CABI)", "TRIP Pro"):
            if not any(name.split(" (")[0].lower() in s.lower()
                       for s in imported):
                self._bullet(f"{name} — not searched")

    # -------------------------------------------------------------- screening

    def _screening_record(self) -> None:
        self._h1("2.  Screening decisions and record counts")

        self._h2("Inclusion criteria")
        for item in self.project.inclusion or ["None recorded."]:
            self._bullet(item)
        self._h2("Exclusion criteria")
        for item in self.project.exclusion or ["None recorded."]:
            self._bullet(item)
        self._para("")
        self._para(f"Minimum level of evidence accepted: {self.project.min_level}",
                   bold=True)

        retrieved = len(self.project.works) + sum(
            1 for n in self.project.excluded_notes if n.get("stage") == "duplicate"
        )
        screened_out = sum(1 for w in self.project.works if w.included is False)
        below_level = sum(
            1 for w in self.project.works
            if w.level.value in ("not-evidence", "ungraded") and w.included is not False
        )
        included = len(self.project.included_works())
        counts = dedupe_module.prisma_counts(
            retrieved=retrieved, after_dedupe=len(self.project.works),
            screened_out=screened_out, excluded_by_level=below_level,
            included=included,
        )

        self._h2("Record counts")
        self._para(
            "These figures are counted from the tool's own search, "
            "deduplication and screening logs. They are the numbers to enter "
            "into a PRISMA 2020 flow diagram; the diagram itself is the property "
            "of the PRISMA group and should be obtained from prisma-statement.org.",
            italic=True, space_after=6)
        self._table(
            ["Stage", "Records"],
            [[k.replace("_", " ").capitalize(), str(v)] for k, v in counts.items()],
            widths=[4.0, 1.2],
        )

        excluded = [w for w in self.project.works if w.included is False]
        if excluded:
            self._para("")
            self._h2("Records excluded at screening, with reasons")
            self._table(
                ["Study", "Reason for exclusion"],
                [[_short(w), w.screen_reason or "No reason recorded"]
                 for w in excluded],
                widths=[2.0, 4.2],
            )

    # --------------------------------------------------------- evidence matrix

    def _evidence_matrix(self) -> None:
        self._landscape()
        self._h1("3.  Evidence matrix of included studies", new_page=False)
        self._para(
            "One row per included study. Level of evidence follows the JBI and "
            "AACN hierarchies; the rationale column records why each study was "
            "graded as it was, so a grade can be checked rather than accepted.",
            italic=True, space_after=8)

        works = sorted(self.project.included_works(),
                       key=lambda w: (w.level.rank, w.sort_key()))
        if not works:
            self._para("No studies are currently included.")
            self._portrait()
            return

        extractions = {e.work_key: e for e in self.project.extractions}
        appraisals = {a.work_key: a for a in self.project.appraisals}

        rows: list[list[str]] = []
        for work in works:
            extraction = extractions.get(work.key)
            appraisal = appraisals.get(work.key)
            rows.append([
                _short(work),
                (extraction.design if extraction else "") or "Not extracted",
                _sample(extraction),
                (extraction.key_findings if extraction else "") or "Not extracted",
                _strengths_limits(extraction, appraisal),
                f"{work.level.value}\n{work.level_reason}"[:220],
            ])
        self._table(
            ["Author(s), year", "Design", "Sample", "Key findings",
             "Strengths and limitations", "Level and rationale"],
            rows, widths=[1.4, 1.2, 1.3, 2.6, 2.3, 1.5],
        )

        self._para("")
        summary = {}
        for work in works:
            summary[work.level.value] = summary.get(work.level.value, 0) + 1
        self._para(
            "Distribution of levels: "
            + ", ".join(f"Level {k} — {v}" for k, v in sorted(summary.items())),
            bold=True)
        self._portrait()

    # -------------------------------------------------------------- appraisal

    def _appraisal_summary(self) -> None:
        self._h1("4.  Critical appraisal of each included study")
        self._para(
            "Each appraisal follows the methodological domains of a published "
            "instrument, named per study. The question wording is this tool's "
            "own rather than the instruments' — CASP, JBI, AGREE II and AMSTAR 2 "
            "are copyrighted, and their item text is not reproduced here. Where a "
            "programme or journal requires a specific completed instrument, "
            "obtain and complete that instrument; this record is a working "
            "appraisal, not a substitute for a form someone else specified.",
            italic=True, space_after=10)

        appraisals = {a.work_key: a for a in self.project.appraisals}
        works = [w for w in self.project.included_works() if w.key in appraisals]
        if not works:
            self._para("No appraisals have been recorded.")
            return

        for work in sorted(works, key=lambda w: w.sort_key()):
            appraisal = appraisals[work.key]
            yes, total = appraisal.score()
            self._h2(_short(work))
            self._para(f"Instrument followed: {appraisal.instrument}", space_after=2)
            self._para(f"Attribution: {appraisal.instrument_citation}",
                       italic=True, space_after=4)
            self._para(
                f"Rating: {appraisal.overall or 'not rated'} — {yes} of {total} "
                f"criteria met. {appraisal.overall_reason}", bold=False, space_after=4)
            unclear = [i for i in appraisal.items if i.answer == "unclear"]
            no = [i for i in appraisal.items if i.answer == "no"]
            if no:
                self._para("Criteria not met:", bold=True, space_after=2)
                for item in no:
                    self._bullet(f"{item.question} {('— ' + item.note) if item.note else ''}")
            if unclear:
                self._para("Criteria unclear from the report as published:",
                           bold=True, space_after=2)
                for item in unclear[:6]:
                    self._bullet(item.question)
                if len(unclear) > 6:
                    self._bullet(f"…and {len(unclear) - 6} more")
            if appraisal.strengths:
                self._para(f"Strengths: {appraisal.strengths}", space_after=2)
            if appraisal.limitations:
                self._para(f"Limitations: {appraisal.limitations}", space_after=8)

    # -------------------------------------------------------------- claim map

    def _claim_map(self) -> None:
        self._landscape()
        self._h1("5.  Claim-to-source traceability map", new_page=False)
        self._para(
            "Every claim in the paper, the source supporting it, where in that "
            "source the support is found, the passage relied on, and why that "
            "source was chosen. Claims marked “author's own analysis” assert "
            "nothing borrowed and therefore carry no citation.",
            italic=True, space_after=8)

        if not self.project.claims:
            self._para("No claims are recorded for this project.")
            self._portrait()
            return

        by_key = {w.key: w for w in self.project.works}
        rows: list[list[str]] = []
        for claim in sorted(self.project.claims, key=lambda c: c.order):
            works = [by_key[k] for k in claim.work_keys if k in by_key]
            if claim.support_type is SupportType.NO_CITATION or not works:
                source = ("Author's own analysis"
                          if claim.support_type is SupportType.NO_CITATION
                          else "⚠ NO SOURCE ATTACHED")
                locus = "—"
            else:
                source = "; ".join(_short(w) for w in works)
                locus = claim.locus or "—"
            rows.append([
                claim.claim_id,
                claim.section,
                claim.text[:420],
                {"paraphrase": "Paraphrase", "direct-quote": "Direct quotation",
                 "data-point": "Data point",
                 "no-citation": "Author's own"}.get(claim.support_type.value, ""),
                source,
                locus,
                (claim.source_quote or "—")[:320],
                (claim.rationale or "—")[:240],
                "yes" if claim.verified else "not yet",
            ])
        self._table(
            ["ID", "Section", "Claim as it appears in the paper", "Use",
             "Source", "Locus", "Passage relied on", "Why this source",
             "Verified"],
            rows,
            widths=[0.5, 1.0, 2.4, 0.7, 1.3, 0.6, 2.0, 1.4, 0.6],
            font_size=self.font_size - 3,
        )
        self._portrait()

    # ------------------------------------------------------- inclusion reasons

    def _exclusion_rationale(self) -> None:
        self._h1("6.  Rationale for inclusion and exclusion")

        self._h2("Why each included study was used")
        works = self.project.cited_works()
        if not works:
            self._para("No sources are cited in the paper yet.")
        else:
            appraisals = {a.work_key: a for a in self.project.appraisals}
            rows = []
            for work in works:
                appraisal = appraisals.get(work.key)
                reasons: list[str] = [work.level.label]
                if appraisal and appraisal.overall:
                    reasons.append(f"appraised {appraisal.overall}")
                extraction = next(
                    (e for e in self.project.extractions if e.work_key == work.key), None)
                if extraction and extraction.sample_size:
                    reasons.append(f"n = {extraction.sample_size}")
                if extraction and extraction.relevance:
                    reasons.append(extraction.relevance)
                if extraction and extraction.setting:
                    reasons.append(f"setting: {extraction.setting}")
                claim_count = sum(1 for c in self.project.claims if work.key in c.work_keys)
                rows.append([
                    _short(work),
                    "; ".join(reasons)[:400],
                    str(claim_count),
                    work.source_db,
                ])
            self._table(
                ["Study", "Basis for inclusion", "Claims supported", "Found via"],
                rows, widths=[1.5, 3.0, 0.8, 1.1])

        self._h2("Studies considered and not used")
        notes = self.project.excluded_notes
        rejected = [w for w in self.project.works if w.included is False]
        if not notes and not rejected:
            self._para(
                "No near-miss exclusions are recorded. Recording these is worth "
                "doing — a reviewer familiar with the topic will ask about the "
                "studies that are missing.")
        else:
            rows = [[n.get("study", ""), n.get("reason", "")] for n in notes]
            rows += [[_short(w), w.screen_reason or "No reason recorded"]
                     for w in rejected]
            self._table(["Study", "Why it was not used"], rows, widths=[2.0, 4.2])

    # -------------------------------------------------------------- integrity

    def _integrity_section(self) -> None:
        self._h1("7.  Retraction and originality checks")

        self._h2("Retraction checks")
        log = [n for n in self.project.excluded_notes if n.get("stage") == "retraction"]
        checked = [w for w in self.project.works if w.included is not False]
        retracted = [w for w in checked if w.retracted]
        self._para(
            f"{len(checked)} source(s) were checked against Crossref retraction "
            f"notices, MEDLINE retraction records and the OpenAlex retraction "
            f"flag. {len(retracted)} retraction(s) were found.", space_after=6)
        if retracted:
            self._table(["Study", "Retraction record"],
                        [[_short(w), w.retraction_note] for w in retracted],
                        widths=[2.0, 4.2])
        for entry in log:
            self._bullet(f"{entry.get('study', '')}: {entry.get('reason', '')}")

        self._h2("Verbatim overlap against cited sources")
        report = integrity_module.check_project(self.project)
        self._para(report.summary(), space_after=8)

        flagged = [o for o in report.overlaps if o.severity in ("serious", "review")]
        if flagged:
            self._table(
                ["Passage", "Words", "Matches", "Severity"],
                [[o.phrase[:200], str(o.word_count), o.work_label, o.severity]
                 for o in flagged[:40]],
                widths=[3.4, 0.6, 1.4, 0.8])

        self._h2("What was not checked")
        guidance = integrity_module.external_scan_guidance([])
        self._para(guidance["interpretation"], space_after=6)
        self._para(
            "No similarity percentage is reported anywhere in this document, "
            "because no defensible target exists. A correct APA paper always "
            "shows measurable similarity: its reference list matches the original "
            "articles exactly, its quotations are supposed to match verbatim, and "
            "the standard phrasing of the field recurs across thousands of "
            "papers. What matters is whether each matched passage is quoted, "
            "cited or reworded — which is what the overlap table above is for.",
            italic=True)

    # ------------------------------------------------------------- disclosure

    def _disclosure(self, drafted: bool) -> None:
        from ..writing.draft import ai_use_disclosure

        self._h1("8.  Disclosure of artificial intelligence use")
        for line in ai_use_disclosure(self.project, drafted).split("\n"):
            if not line.strip():
                self._para(space_after=4)
            elif line.endswith("use") and line.startswith("Disclosure"):
                continue  # the heading above already says this
            else:
                self._para(line, space_after=4)

    # -------------------------------------------------------------- references

    def _references(self) -> None:
        self._h1("9.  References")
        self._para(
            "Identical to the paper's reference list. Reproduced here so this "
            "document can be read on its own.", italic=True, space_after=8)
        works = self.project.cited_works()
        if not works:
            self._para("No sources are cited.")
            return
        for entry in reference_list(works, self.context):
            paragraph = self._para(runs=entry, space_after=6)
            fmt = paragraph.paragraph_format
            fmt.left_indent = Inches(0.5)
            fmt.first_line_indent = Inches(-0.5)

    # ------------------------------------------------------------------- save

    def save(self, path: str) -> str:
        self.document.save(path)
        return path


# ------------------------------------------------------------------ helpers


def _short(work: Work) -> str:
    surname = work.first_author_surname() or "Anonymous"
    others = len(work.authors or work.editors)
    if others > 2:
        surname += " et al."
    elif others == 2:
        second = (work.authors or work.editors)[1].intext_name()
        surname += f" & {second}"
    return f"{surname} ({work.year or 'n.d.'})"


def _sample(extraction) -> str:
    if not extraction:
        return "Not extracted"
    parts = [p for p in (extraction.sample_description, extraction.setting) if p]
    text = "; ".join(parts)
    if extraction.sample_size:
        text = f"n = {extraction.sample_size}" + (f"; {text}" if text else "")
    return text or "Not extracted"


def _strengths_limits(extraction, appraisal) -> str:
    parts: list[str] = []
    if extraction and extraction.strengths:
        parts.append(f"Strengths: {extraction.strengths}")
    if extraction and extraction.limitations:
        parts.append(f"Limitations: {extraction.limitations}")
    if appraisal and appraisal.limitations and not (extraction and extraction.limitations):
        parts.append(f"Appraisal: {appraisal.limitations}")
    return " ".join(parts)[:400] or "Not extracted"


def build_audit_document(
    project: Project, context: CitationContext, path: str, *, drafted: bool = True,
    font: str = "Times New Roman", font_size: float = 12.0,
) -> str:
    return AuditDocument(project, context, font, font_size).build(
        drafted=drafted).save(path)
