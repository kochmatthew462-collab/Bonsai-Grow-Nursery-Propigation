"""
The parts of a Word document python-docx has no API for.

APA 7 needs four things that live below python-docx's surface, and each one is
here because there is no supported way to do it otherwise:

* **Automatic page numbers** are a Word *field*, not text. python-docx can
  write a header paragraph but cannot insert `{ PAGE }`, so the field's three
  OOXML runs are built by hand. Writing the literal digit instead — which is
  what most generated documents do — produces a paper where every page is
  numbered 1.
* **Font names** have four independent attributes in OOXML (`w:ascii`,
  `w:hAnsi`, `w:eastAsia`, `w:cs`). `style.font.name` sets one of them, so a
  document with an em dash or a Greek letter silently falls back to a
  different typeface mid-line.
* **APA tables have horizontal rules only** — no vertical borders, no grid.
  Table borders are a property of the table's XML, not of any python-docx
  object.
* **Widow/orphan control and keep-with-next** on headings, so a Level 2
  heading never sits alone at the foot of a page.

Everything here takes a python-docx object and mutates its `._element`. That
private attribute is the documented escape hatch for exactly this purpose.
"""

from __future__ import annotations

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt


# ------------------------------------------------------------------- fields


def add_page_number_field(paragraph) -> None:
    """Insert a `{ PAGE }` field so Word numbers the pages itself.

    A field is three runs: a `begin` fldChar, the instruction text, and an
    `end` fldChar. Word evaluates it on open and on print.
    """
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"

    # A cached result, so the number is visible even in readers that do not
    # evaluate fields (Google Docs preview, some PDF converters).
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    cached = OxmlElement("w:t")
    cached.text = "1"

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    for element in (begin, instruction, separate, cached, end):
        run._r.append(element)


def add_total_pages_field(paragraph) -> None:
    """Insert `{ NUMPAGES }`. Not required by APA, but used by the audit
    document's "page x of y" footer."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "NUMPAGES"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, end):
        run._r.append(element)


# --------------------------------------------------------------------- fonts


def force_font(run, name: str, size_pt: float | None = None) -> None:
    """Set all four OOXML font attributes on one run.

    `w:ascii` alone covers only Latin-1. Without `w:hAnsi` an en dash can shift
    typeface, and without `w:cs` (complex script) so can anything from a
    right-to-left range. APA's font requirement is about the whole document,
    not just its ASCII.
    """
    run.font.name = name
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for attribute in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(attribute), name)


def force_style_font(style, name: str, size_pt: float) -> None:
    """Same, applied to a style rather than a run."""
    style.font.name = name
    style.font.size = Pt(size_pt)
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for attribute in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(attribute), name)


# ------------------------------------------------------------ paragraph flags


def keep_with_next(paragraph, on: bool = True) -> None:
    """Stop a heading being orphaned at the foot of a page."""
    ppr = paragraph._p.get_or_add_pPr()
    _set_flag(ppr, "w:keepNext", on)


def keep_lines_together(paragraph, on: bool = True) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    _set_flag(ppr, "w:keepLines", on)


def widow_control(paragraph, on: bool = True) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    _set_flag(ppr, "w:widowControl", on)


def suppress_auto_hyphenation(paragraph) -> None:
    """APA 7 §2.23: do not hyphenate; leave the right margin ragged."""
    ppr = paragraph._p.get_or_add_pPr()
    _set_flag(ppr, "w:suppressAutoHyphens", True)


def _set_flag(ppr, tag: str, on: bool) -> None:
    element = ppr.find(qn(tag))
    if element is None:
        element = OxmlElement(tag)
        ppr.append(element)
    element.set(qn("w:val"), "1" if on else "0")


def page_break_before(paragraph) -> None:
    """A page break as a paragraph property rather than a break character.

    This is what keeps "References" at the top of its own page even after the
    body text is edited — a manual break run would drift.
    """
    ppr = paragraph._p.get_or_add_pPr()
    element = ppr.find(qn("w:pageBreakBefore"))
    if element is None:
        element = OxmlElement("w:pageBreakBefore")
        ppr.append(element)
    element.set(qn("w:val"), "1")


# --------------------------------------------------------------------- tables


def apa_table_borders(table) -> None:
    """APA 7 §7.8-7.21: horizontal rules only.

    A rule above and below the column headings, one at the foot of the table,
    and nothing else — no vertical lines and no interior horizontal lines
    between data rows. python-docx's built-in "Table Grid" style is the
    opposite of this on every count.
    """
    tbl_pr = table._tbl.tblPr
    for existing in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(existing)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")          # half-points: a 0.5 pt rule
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        borders.append(element)
    for edge in ("left", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "none")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        borders.append(element)

    # `w:tblPr` children are order-sensitive: the schema fixes the sequence, and
    # `w:tblBorders` must precede `w:shd`, `w:tblLayout`, `w:tblCellMar` and
    # `w:tblLook`. python-docx's default table style already emits `w:tblLook`,
    # so a plain append lands the borders *after* it and produces a document
    # Word may repair silently — dropping the borders and restoring the grid,
    # which is the opposite of an APA table.
    tbl_pr.insert_element_before(
        borders, "w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook",
    )

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _repeat_header_row(table)


def rule_below_row(row) -> None:
    """A single rule under one row — used under the column headings."""
    for cell in row.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        bottom = borders.find(qn("w:bottom"))
        if bottom is None:
            bottom = OxmlElement("w:bottom")
            borders.append(bottom)
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "0")
        bottom.set(qn("w:color"), "000000")


def _repeat_header_row(table) -> None:
    """Repeat the heading row when a table spans pages (APA 7 §7.13)."""
    if not table.rows:
        return
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def clear_cell_margins_shading(table) -> None:
    """APA tables carry no shading. Some Word templates add it by default."""
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            for shd in tc_pr.findall(qn("w:shd")):
                tc_pr.remove(shd)


# -------------------------------------------------------------------- headers


def build_apa_header(section, running_head: str = "") -> None:
    """The APA 7 page header.

    Professional papers carry a running head flush left in capitals plus the
    page number flush right; student papers carry the page number alone
    (APA 7 §2.4, §2.8, §2.18). Both appear on every page, the title page
    included — which is why `different_first_page_header_footer` is left off.
    """
    section.different_first_page_header_footer = False
    header = section.header
    header.is_linked_to_previous = False

    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.text = ""
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)

    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0

    if running_head:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_right_tab(paragraph, section)
        paragraph.add_run(running_head.upper()[:50])
        paragraph.add_run("\t")
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    add_page_number_field(paragraph)


def _set_right_tab(paragraph, section) -> None:
    """A right-aligned tab stop at the right margin, so the page number lands
    flush right on the same line as the running head.

    The unit conversion is the whole trick and it is easy to get wrong: OOXML tab
    positions are in twentieths of a point (twips), and there are exactly 635 EMU
    to a twip. An earlier version multiplied the result by ten on the assumption
    that `w:pos` was in some finer unit, which put the tab stop at 60 inches —
    far off the page, so the page number wrapped to a second header line instead
    of sitting beside the running head. 6.5 inches of text width is 9360 twips.
    """
    width_twips = int(
        (section.page_width - section.left_margin - section.right_margin) / 635
    )
    ppr = paragraph._p.get_or_add_pPr()
    tabs = ppr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        ppr.append(tabs)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(width_twips))
    tabs.append(tab)


def set_document_language(document, language: str = "en-US") -> None:
    """Tag the document language so Word's own proofing tools engage.

    This matters for the grammar workflow: the suite proofs text before export,
    but the user will also run Word's or Grammarly's checker over the finished
    file, and an untagged document gets checked against the wrong dictionary.
    """
    styles = document.styles.element
    rpr = styles.find(qn("w:docDefaults"))
    normal = document.styles["Normal"].element.get_or_add_rPr()
    lang = normal.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        normal.append(lang)
    lang.set(qn("w:val"), language)
    lang.set(qn("w:eastAsia"), language)
    lang.set(qn("w:bidi"), "ar-SA")
    del rpr  # docDefaults is looked up only to assert the styles part exists
